import docx
from fastapi import APIRouter, HTTPException, Depends, Request, Form, Query
from sqlalchemy.orm import Session, joinedload 
from typing import List, Tuple, Optional, Dict, Union
from ..model.database.database import SessionLocal
from ..model.pacient_hospital import models, schemas, old_models
from ..model.сlinic.models import DoctorsOfTheClinic 
from datetime import datetime, date, timedelta
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from sqlalchemy import func, and_, or_
from fastapi.responses import JSONResponse
import pandas as pd
from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.worksheet.page import PageMargins
from openpyxl.worksheet.worksheet import Worksheet
from docx import Document
from starlette.responses import FileResponse
from docx.shared import Inches
import io
import os
from urllib.parse import quote
from pytz import timezone
from io import BytesIO
from sqlalchemy.orm import aliased
from starlette.middleware.base import BaseHTTPMiddleware
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
    
router = APIRouter()

frontend_dir = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "../templates"))

current_dir = os.path.dirname(os.path.abspath(__file__))
docs_dir = os.path.abspath(os.path.join(current_dir, "../static/docs"))

templates = Jinja2Templates(directory=frontend_dir)

from ._utils.check_utils import (format_datetime, 
                                 get_moscow_time,
                                 remove_tzinfo_datetime,
                                 remove_tzinfo)

templates.env.filters["format_datetime"] = format_datetime

MIN_CHAR_COUNT = 3  # Минимальное количество символов для поля, чтобы оно считалось заполненным

# Список символов, которые будут игнорироваться при проверке заполненности полей
EXCLUDED_CHARACTERS = [' ', '-', '_', '.']


def remove_excluded_characters(value: str) -> str:
    """
    Удаляет исключенные символы из строки.

    Args:
        value (str): Входная строка.

    Returns:
        str: Строка без исключенных символов.
    """
    for char in EXCLUDED_CHARACTERS:
        value = value.replace(char, '')
    return value


def check_field_completeness(value: Union[str, date], field_name: str, min_char_count: int = MIN_CHAR_COUNT) -> bool:
    """
    Проверяет заполненность поля. Для строк проверяет количество символов, для дат — всегда считается заполненным.

    Args:
        value (Union[str, date]): Значение поля.
        field_name (str): Имя поля.
        min_char_count (int): Минимальное количество символов для строки, чтобы поле считалось заполненным.

    Returns:
        bool: True, если поле считается заполненным, иначе False.
    """
    if isinstance(value, date):
        return True  # Дата всегда считается заполненной
    elif isinstance(value, str):
        value = remove_excluded_characters(value)
        return len(value) >= min_char_count
    else:
        return False

def check_patient_data_completeness(passport_data: models.PassportData) -> Tuple[List[str], float]:
    incomplete_fields = []

    # Поля для PassportData
    passport_fields = [
        'full_name', 'birthday_date', 'personal_data', 'military_rank',
        'directions', 'date_of_illness', 'address', 'military_unit',
        'history_number', 'phone_number', 'branch', 'service_basis',
        'personal_document'
    ]

    # Проверка полей модели PassportData
    for field in passport_fields:
        value = getattr(passport_data, field, None)
        if not check_field_completeness(value, field):
            incomplete_fields.append(field)

    # Поля для HospitalData
    hospital_fields = [
        'vk_urgent_call_date', 'vk_call_up_date', 'district',
        'diagnosis_upon_admission', 'final_diagnosis', 'ICD',
        'character_of_the_hospital', 'reason_for_departure',
        'vvk_decision', 'anamnesis', 'expert_diagnosis',
        'diagnosis_according_to_form_one_hundred', 'therapist'
    ]

    # Проверка связанных госпитальных данных
    for hospital_data in passport_data.hospital_records:
        for field in hospital_fields:
            value = getattr(hospital_data, field, None)
            if not check_field_completeness(value, field):
                incomplete_fields.append(f"{field} (hospital_record)")

    # Общий список полей для расчета процента 
    total_fields_count = len(passport_fields) + len(hospital_fields) * len(passport_data.hospital_records)
    filled_fields_count = total_fields_count - len(incomplete_fields)
    
    completeness_percentage = (filled_fields_count / total_fields_count) * 100 if total_fields_count else 100


    return incomplete_fields, completeness_percentage



# Функция для получения сессии базы данных
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

#region old_data
@router.get("/hospitalized_patients_old", response_class=HTMLResponse)
async def read_root(request: Request, db: Session = Depends(get_db)):
    # Fetching data from the database
    patients = db.query(old_models.PatientData).all()

    # Prepare data for rendering
    data = {
        'request': request,
        'patients': [{
            'history_number': patient.history_number,
            'military_rank': patient.military_rank,
            'service_basis': patient.service_basis,
            'military_unit': patient.military_unit,
            'military_commissariat': patient.military_commissariat,
            'district': patient.district,
            'full_name': patient.full_name,
            'birthday_date': patient.birthday_date,
            'address': patient.address,
            'phone_number': patient.phone_number,
            'personal_data': patient.personal_data,
            'availability': patient.availability,
            'assigned_by': patient.assigned_by,
            'admission_date': patient.admission_date,
            'illness_date': patient.illness_date,
            'diagnosis': patient.diagnosis,
            'department': patient.department,
            'nature_of_hospitalization': patient.nature_of_hospitalization,
            'transfer': patient.transfer,
            'departure': patient.departure
        } for patient in patients]
    }

    return templates.TemplateResponse("old_patient/patient_table.html", data)
#endregion


# region trash

# Разрешенные IP-адреса admin
ALLOWED_IPS = ["127.0.0.1","192.168.15.99", "192.168.15.205", "192.168.15.203", "192.168.15.228", "192.168.15.115", "192.168.15.148", "192.168.15.114", "192.168.15.192"]  # Замените на разрешенные IP

# Разрешенные IP-адреса смены ВВК и Справок о тяжести
ALLOWED_IPS_VVK = ["127.0.0.1", "192.168.15.205", "192.168.15.203", "192.168.15.228", "192.168.15.115", "192.168.15.76"] 

# Разрешенные IP-vvk

def check_ip_vvk(request: Request):
    client_ip = request.client.host
    if client_ip not in ALLOWED_IPS_VVK:
        raise HTTPException(status_code=403, detail="Доступ запрещен.")

# Зависимость для проверки IP-адреса
def check_ip(request: Request):
    client_ip = request.client.host
    if client_ip not in ALLOWED_IPS:
        raise HTTPException(status_code=403, detail="Доступ запрещен.")






# endregion

# region excel low


# Для модели PassportData
passport_data_columns = {
    'full_name': 'ФИО',
    'birthday_date': 'Дата рождения',
    'personal_data': 'Личный номер',
    'military_rank': 'Воинское звание',
    'current_time': 'Время госпитализации',
    'directions': 'Направление',
    'date_of_illness': 'Дата заболевания',
    'address': 'Адрес',
    'military_unit': 'Воинская часть',
    'history_number': 'Номер истории',
    'phone_number': 'Номер телефона',
    'branch': 'Отделение',
    'service_basis': 'Основа службы',
    'personal_document': 'Документ (Военный билет/паспорт)',
    'nature_of_hospitalization': 'Характер госпитализации',
    'first_diagnosis': "Начальный диагноз",
    'after_hostilities': "После боевых действий (СВО)"
}
    
# Для модели PatientMovement
patient_movement_columns = {
    'department': 'Отделение',
    'event_type': 'Тип события',
    'event_date': 'Дата события',
    'destination_department': 'Целевое отделение',
    'note': 'Примечание'
}

# Для модели HospitalData
hospital_data_columns = {
    'vk_urgent_call_date': 'Дата срочного вызова ВК',
    'vk_call_up_date': 'Дата вызова ВК',
    'district': 'Район',
    'diagnosis_upon_admission': 'Диагноз при поступлении',
    'final_diagnosis': 'Окончательный диагноз',
    'ICD': 'МКБ',
    'character_of_the_hospital': 'Характер больницы',
    'reason_for_departure': 'Причина отбытия',
    'vvk_decision': 'Решение ВВК',
    'certificate_of_injury': 'Справка о травме',
    'medical_record': 'Медицинская карта',
    'food_certificate': 'Справка о питании',
    'sick_leave': 'Больничный лист',
    'entered_after_participating_in_hostilities': 'Поступление после участия в боевых действиях',
    'suitability_category': 'Не достиг части',
    'anamnesis': 'Анамнез',
    'severity_of_injury': 'Тяжесть ранения',
    'expert_diagnosis': 'Экспертный диагноз',
    'diagnosis_according_to_form_one_hundred': 'Диагноз по справке 100',
    'therapist': 'Лечящий врач'
}

certificate_of_injury_columns = {
    'certificate_injury_date': "Дата выдачи справки о ранении",
    'injury_number': 'Номер справки о ранении',
    'load_date': "Дата загрузки в Алушту"
}

certificate_of_severity_columns = {
    "severity_number": "Номер справки о тяжести",
    'severity_date': 'Дата выдачи справки о тяжести',
    'approval_date': 'Дата подтверждения справки о тяжести',
    'approval_number': 'Номер подтверждения справки о тяжести'
}


from ._utils.excel_utils import (apply_filters_to_excel, adjust_column_widths, get_latest_patient_movements)

# Функция для удаления временной зоны и преобразования в строки только с датой для всех остальных столбцов


@router.get("/export_data_to_excel")
async def export_data_to_excel(db: Session = Depends(get_db)):  
    """
    Экспорт данных о пациентах в Excel с статичным размером столбцов.

    Эта функция извлекает данные из таблиц PassportData, PatientMovement, HospitalData, CertificateOfInjury и CertificateOfSeverity,
    удаляет информацию о временной зоне из всех столбцов datetime, объединяет данные в один DataFrame,
    сохраняет результат в Excel и возвращает файл для скачивания. Ширина столбцов настраивается
    в зависимости от длины названий столбцов.

    Args:
        db (Session): Сессия базы данных, передаваемая через Depends.

    Returns:
        StreamingResponse: Файл Excel с объединенными данными о пациентах.
    """
    try:
        # Извлечение данных из таблиц базы данных
        passport_data_df = pd.read_sql(db.query(models.PassportData).statement, db.bind)
        latest_patient_movements_data = get_latest_patient_movements(db)
        latest_patient_movements_df = pd.DataFrame(latest_patient_movements_data)
        hospital_data_df = pd.read_sql(db.query(models.HospitalData).statement, db.bind)
        certificate_of_injury_df = pd.read_sql(db.query(models.CertificateOfInjury).statement, db.bind)
        certificate_of_severity_df = pd.read_sql(db.query(models.CertificateOfSeverity).statement, db.bind)

        # Удаление информации о временной зоне
        passport_data_df = remove_tzinfo_datetime(passport_data_df)
        latest_patient_movements_df = remove_tzinfo(latest_patient_movements_df)
        hospital_data_df = remove_tzinfo(hospital_data_df)
        certificate_of_injury_df = remove_tzinfo(certificate_of_injury_df)
        certificate_of_severity_df = remove_tzinfo(certificate_of_severity_df)

        # Переименование столбцов (необходимо определить словари переименования)
        passport_data_df.rename(columns=passport_data_columns, inplace=True)
        latest_patient_movements_df.rename(columns=patient_movement_columns, inplace=True)
        hospital_data_df.rename(columns=hospital_data_columns, inplace=True)
        certificate_of_injury_df.rename(columns=certificate_of_injury_columns, inplace=True)
        certificate_of_severity_df.rename(columns=certificate_of_severity_columns, inplace=True)

        # Преобразование datetime столбцов в строку
        for df in [passport_data_df, latest_patient_movements_df, hospital_data_df, certificate_of_injury_df, certificate_of_severity_df]:
            for col in df.select_dtypes(include=['datetime']):
                df[col] = df[col].astype(str)

        # Объединение всех DataFrame
        combined_df = passport_data_df.merge(latest_patient_movements_df, left_on='id', right_on='patient_id', how='left', suffixes=('', '_movement'))
        combined_df = combined_df.merge(hospital_data_df, on='patient_id', how='left', suffixes=('', '_hospital'))
        combined_df = combined_df.merge(certificate_of_injury_df, on='patient_id', how='left', suffixes=('', '_injury'))
        combined_df = combined_df.merge(certificate_of_severity_df, on='patient_id', how='left', suffixes=('', '_severity'))

        # Удаление ненужных столбцов
        columns_to_drop = ['id', 'id_movement', 'id_hospital', 'id_injury', 'id_severity']
        combined_df.drop(columns=[col for col in columns_to_drop if col in combined_df.columns], inplace=True)




        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            combined_df.to_excel(writer, sheet_name='Архив', index=False)
            
            # Настройка фиксированной ширины столбцов
            worksheet = writer.sheets['Архив']
            cm_to_units = 3 / 0.14  # 3 см, конвертированные в единицы Excel
            fixed_width = int(cm_to_units)
            for col in worksheet.columns:
                column = col[0].column_letter
                worksheet.column_dimensions[column].width = fixed_width

        output.seek(0)  # Сброс позиции буфера в начало для ответа

        output_with_filters = apply_filters_to_excel(output)

        current_datetime = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"Учёт госпитализированных на {current_datetime}.xlsx"
        safe_filename = quote(filename)

        return StreamingResponse(output_with_filters, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                 headers={'Content-Disposition': f'attachment; filename="{safe_filename}"'})

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def update_certificate_of_severity(db: Session, record_id: int, updated_data: schemas.CertificateOfSeverityUpdate):
    db_record = db.query(models.CertificateOfSeverity).filter(models.CertificateOfSeverity.id == record_id).first()
    if db_record:
        for key, value in updated_data.dict().items():
            setattr(db_record, key, value)
        db.commit()
        db.refresh(db_record)
        return db_record
    else:
        return None

def get_certificate_of_severity(db: Session, record_id: int):
    return db.query(models.CertificateOfSeverity).filter(models.CertificateOfSeverity.id == record_id).first()

def get_certificate_of_injury(db: Session, record_id: int):
    return db.query(models.CertificateOfInjury).filter(models.CertificateOfInjury.id == record_id).first()

def update_certificate_of_injury(db: Session, record_id: int, updated_data: dict):
    db_record = db.query(models.CertificateOfInjury).filter(models.CertificateOfInjury.id == record_id).first()
    if db_record:
        for key, value in updated_data.items():
            setattr(db_record, key, value)
        db.commit()
        db.refresh(db_record)
        return db_record
    else:
        return None



#endregion

# region sss


@router.get("/patients_with_pneumonia/")
async def get_patients_with_pneumonia(request: Request, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    
    patients_with_pneumonia = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .filter(models.HospitalData.diagnosis_upon_admission.ilike("%пневмония%"))  # Фильтрация по диагнозу
        .order_by(models.PassportData.current_time.desc())
        .all()
    )


    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_with_pneumonia:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Пациенты, с диагнозом %пневмония%",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

@router.get("/passport_data_active/")
async def filtered_passport_data_html(request: Request, db: Session = Depends(get_db)):
    """
    Получение данных активных пациентов для отображения в HTML-шаблоне.

    Эта функция извлекает данные о пациентах, которые были приняты на госпитализацию или переведены
    в другое отделение, и обрабатывает их для отображения в HTML-таблице. Также рассчитывается
    процент заполненности данных пациента и определяется цветовой код для визуализации
    этого процента. Данные о дате выписки пациента также извлекаются, если пациент был выписан.

    Args:
        request (Request): Запрос для рендеринга HTML-шаблона.
        db (Session): Сессия базы данных, передаваемая через Depends.

    Returns:
        TemplateResponse: HTML-шаблон с данными о пациентах.
    """
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            models.PatientMovement.event_type.in_([
                "Прием на госпитализацию",
                "Перевод в другое отделение",
                "Планируется перевод"
            ])
        ))
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = False  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Пациенты, находящиеся на госпитализации",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

@router.get("/transfer_document/")
async def create_transfer_document(date: str = None, db: Session = Depends(get_db)):
    """
    Создание документа с таблицей пациентов на перевод.

    Args:
        date (str): Дата для отображения в заголовке документа.
        db (Session): Сессия базы данных, передаваемая через Depends.

    Returns:
        StreamingResponse: Ответ с документом для скачивания.
    """
    if date:
        selected_date = datetime.strptime(date, "%Y-%m-%d")
    else:
        selected_date = datetime.now()

    # Извлечение данных пациентов
    start_of_day = selected_date.replace(hour=0, minute=0, second=0, microsecond=0)
    end_of_day = start_of_day + timedelta(days=1)


    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData, models.PatientMovement.note)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == models.PassportData.id,
            models.PatientMovement.event_type == "Планируется перевод",
            models.PatientMovement.event_date >= start_of_day,
            models.PatientMovement.event_date < end_of_day
        ))
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

   

    # Создание документа
    doc = Document()

    # Установка ориентации страницы на альбомную
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width, section.page_height = new_width, new_height

    # Установка полей страницы
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Установка шрифта и размера для всего документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Добавление заголовка
    header = doc.add_heading(level=1)
    header_run = header.add_run('ВОЕННЫЙ ГОСПИТАЛЬ (на 150 коек, г. Казань) ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ КАЗЕННОЕ УЧРЕЖДЕНИЕ «354 ВОЕННЫЙ КЛИНИЧЕСКИЙ ГОСПИТАЛЬ» МО РФ')
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(16)
    header_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subheader = doc.add_heading(level=2)
    subheader_run = subheader.add_run(f'Пациенты на перевод на {selected_date.strftime("%d.%m.%Y")} на 14-00 для стационарного лечения.')
    subheader_run.font.name = 'Times New Roman'
    subheader_run.font.size = Pt(14)
    subheader_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Создание таблицы
    table = doc.add_table(rows=1, cols=11)  # Изменение количества столбцов на 9
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['№', 'ФИО', 'Дата рождения', 'Основа службы', 'Округ', 'Звание', 'Тип службы', 'Окончательный диагноз', 'Код МКБ', 'Личный номер','Примечание']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    # Заполнение таблицы данными
    for index, (passport_data, hospital_data, note) in enumerate(patients_admitted_on_date, start=1):
        row_cells = table.add_row().cells
        data = [
            str(index),  # Добавление порядкового номера
            passport_data.full_name,
            passport_data.birthday_date.strftime("%d.%m.%Y") if passport_data.birthday_date else '',
            passport_data.service_basis if passport_data.service_basis else '',
            hospital_data.district if hospital_data.district else '',
            passport_data.military_rank if passport_data.military_rank else '',
            passport_data.branch,
            hospital_data.final_diagnosis if hospital_data.final_diagnosis else '',
            hospital_data.ICD if hospital_data.ICD else '',
            passport_data.personal_data if passport_data.personal_data else '',
            note if note else ''
        ]
        for i, cell_data in enumerate(data):
            row_cells[i].text = cell_data
            run = row_cells[i].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    # Сохранение документа в памяти
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    # Кодирование имени файла для правильной передачи в заголовке Content-Disposition
    import urllib.parse
    filename = f"Пациенты на перевод на {selected_date.strftime('%Y-%m-%d')}.docx"
    encoded_filename = urllib.parse.quote(filename)

    # Возвращение документа как ответа
    return StreamingResponse(byte_io, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
    })

@router.get("/passport_data_planned/")
async def filtered_passport_data_html(request: Request, date: str = None, db: Session = Depends(get_db)):
    """
    Получение данных активных пациентов для отображения в HTML-шаблоне.

    Эта функция извлекает данные о пациентах, которые были приняты на госпитализацию или переведены
    в другое отделение, и обрабатывает их для отображения в HTML-таблице. Также рассчитывается
    процент заполненности данных пациента и определяется цветовой код для визуализации
    этого процента. Данные о дате выписки пациента также извлекаются, если пациент был выписан.

    Args:
        request (Request): Запрос для рендеринга HTML-шаблона.
        db (Session): Сессия базы данных, передаваемая через Depends.

    Returns:
        TemplateResponse: HTML-шаблон с данными о пациентах.
    """
    if date:
        selected_date = datetime.strptime(date, "%Y-%m-%d")
    else:
        selected_date = datetime.now()

    start_of_day = selected_date.replace(hour=0, minute=0, second=0, microsecond=0)
    end_of_day = start_of_day + timedelta(days=1)

   
    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == models.PassportData.id,
            models.PatientMovement.event_type == "Планируется перевод",
            models.PatientMovement.event_date >= start_of_day,
            models.PatientMovement.event_date < end_of_day
        ))
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )


    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": f"Планируется перевод на {selected_date.strftime('%d.%m.%Y')}",
        "selected_date": selected_date.strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })


@router.get("/hospitalized_patients/", response_class=HTMLResponse)
async def hospitalized_patients(request: Request, db: Session = Depends(get_db)):
    # Получаем общее количество пациентов
    total_patients = db.query(models.PassportData).count()  # Подсчет количества пациентов
    return templates.TemplateResponse("patient_hospital/patient_table_archive.html", {"request": request, 
                                                                                      "total_patients": total_patients,
                                                                                      "title": "Архив пациентов"})

@router.get("/hospitalized_patients/search/", response_class=JSONResponse)
async def search_hospitalized_patients(query: str, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для поиска пациентов по ФИО в формате JSON.

    Параметры:
        query (str): Строка для поиска по ФИО пациентов.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        JSONResponse: Данные пациентов в формате JSON, соответствующие запросу.
    """
    # Запрос для поиска пациентов по ФИО
    patients = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .filter(models.PassportData.full_name.ilike(f"%{query}%"))
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date_row = (
            db.query(models.PatientMovement.event_date)
            .filter(
                models.PatientMovement.patient_id == passport_data.id,
                models.PatientMovement.event_type.in_(["Выписан", 
                                                       "Выписан за нарушение режима",
                                                       "Перевод в другое МО",
                                                       "Перевод в другое ВМО",
                                                       "Направлен в санаторий",
                                                       "СОЧ",
                                                       "Летальный исход"])
            )
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )
        discharge_date = format_datetime(discharge_date_row.event_date) if discharge_date_row else None

        # Добавление обработанных данных в список
        patient_data_with_completeness.append({
            "id": passport_data.id,
            "history_number": passport_data.history_number,
            "full_name": passport_data.full_name,
            "birthday_date": passport_data.birthday_date.strftime('%d.%m.%Y') if passport_data.birthday_date else '',
            "military_rank": passport_data.military_rank,
            "military_unit": passport_data.military_unit or '',
            "service_basis": passport_data.service_basis,
            "district": hospital_data.district or '',
            "directions": passport_data.directions or '',
            "branch": passport_data.branch,
            "current_time": passport_data.current_time.strftime('%d.%m.%Y в %H:%M'),
            "suitability_category": hospital_data.suitability_category or '',
            "first_diagnosis": passport_data.first_diagnosis or '',
            "diagnosis_upon_admission": hospital_data.diagnosis_upon_admission or '',
            "final_diagnosis": hospital_data.final_diagnosis or '',
            "therapist": hospital_data.therapist or '',
            "discharge_date": discharge_date or '',
            "completeness_percentage": completeness_percentage,
            "color": color
        })

    return JSONResponse(content=patient_data_with_completeness)

@router.get("/load_patients_data/", response_class=HTMLResponse)
async def load_patients_data(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов вместе с их
    госпитальными записями, процентом заполненности и, при необходимости, датами выписки.

    Параметры:
        request (Request): Объект запроса.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        TemplateResponse: Отображает данные пациентов в HTML-шаблоне.
    """
    patients_admitted_on_date: List[Tuple[models.PassportData, models.HospitalData]] = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        discharge_date_row = (
            db.query(models.PatientMovement.event_date)
            .filter(
                models.PatientMovement.patient_id == passport_data.id,
                models.PatientMovement.event_type.in_(["Выписан", 
                                                       "Выписан за нарушение режима",
                                                        "Перевод в другое МО",
                                                        "Перевод в другое ВМО",
                                                        "Направлен в санаторий",
                                                        "СОЧ",
                                                        "Летальный исход"])
            )
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )
        discharge_date = format_datetime(discharge_date_row.event_date) if discharge_date_row else None

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    return templates.TemplateResponse("patient_hospital/patient_table_archive.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Архив пациентов",
        "hide_discharge_date": False
    })

@router.get("/discharged_on_date/")
async def get_discharged_on_date(request: Request, date: str = None, db: Session = Depends(get_db)):
    if date:
        selected_date = datetime.strptime(date, "%Y-%m-%d")
    else:
        selected_date = datetime.now()

    start_of_day = selected_date.replace(hour=0, minute=0, second=0, microsecond=0)
    end_of_day = start_of_day + timedelta(days=1)

    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).filter(
        models.PatientMovement.event_type.in_([
            "Выписан",
            "Выписан за нарушение режима",
            "Перевод в другое ВМО",
            "Перевод в другое МО",
            "Направлен в санаторий",
            "СОЧ",
            "Летальный исход",
            "Перевод в РКБ"
        ])
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_discharged_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Выписан за нарушение режима",
                "Перевод в другое ВМО",
                "Перевод в другое МО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ"
            ]),
            models.PatientMovement.event_date >= start_of_day,
            models.PatientMovement.event_date < end_of_day
        ))
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_discharged_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": f"Пациенты, выписанные за {selected_date.strftime('%d.%m.%Y')}",
        "selected_date": selected_date.strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

@router.get("/rkb/")
async def get_discharged_on_date(request: Request, date: str = None, db: Session = Depends(get_db)):
    if date:
        selected_date = datetime.strptime(date, "%Y-%m-%d")
    else:
        selected_date = datetime.now()

    start_of_day = selected_date.replace(hour=0, minute=0, second=0, microsecond=0)
    end_of_day = start_of_day + timedelta(days=1)

    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_discharged_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            models.PatientMovement.event_type == "Перевод в РКБ",
            models.PatientMovement.event_date >= start_of_day,
            models.PatientMovement.event_date < end_of_day
        ))
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_discharged_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": f"Пациенты, выписанные за {selected_date.strftime('%d.%m.%Y')}",
        "selected_date": selected_date.strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

    
#endregion

# region discharged at branch

@router.get("/hospitalized_patients_data_on_date/")
async def passport_data_on_date(request: Request, date: str = None, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов, принятых на госпитализацию
    в определенный день. Данные включают процент заполненности и цветовой код, а также опционально
    могут включать дату выписки.

    Параметры:
        request (Request): Объект запроса.
        date (str, optional): Строка даты в формате "ГГГГ-ММ-ДД". Если не указана, используется текущая дата.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        TemplateResponse: Отображает данные пациентов в HTML-шаблоне.
    """
    if date:
        selected_date = datetime.strptime(date, "%Y-%m-%d")
    else:
        selected_date = datetime.now()

    start_of_day = selected_date.replace(hour=0, minute=0, second=0, microsecond=0)
    end_of_day = start_of_day + timedelta(days=1)

    subquery = db.query(
        models.PatientMovement.patient_id,
        func.min(models.PatientMovement.event_date).label('earliest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.earliest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение" 
            ),
            models.PatientMovement.event_date >= start_of_day,
            models.PatientMovement.event_date < end_of_day
        ))
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": f"Пациенты, принятые на госпитализацию {selected_date.strftime('%d.%m.%Y')}",
        "selected_date": selected_date.strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })


@router.get("/hospitalized_patients_infectious_diseases_department/")
async def hospitalized_patients_infectious_diseases_department(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов из неврологического отделения,
    которые были выписаны или переведены по различным причинам.

    Параметры:
        request (Request): Объект запроса.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        TemplateResponse: Отображает данные пациентов в HTML-шаблоне.
    """
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима" 
            ])
        ))
        .filter(models.PassportData.branch == "Инфекционное отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = False  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Выписанные пациенты Инфекционного отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
  
@router.get("/hospitalized_patients_neurological_department/")
async def hospitalized_patients_neurological_department(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов из неврологического отделения,
    которые были выписаны или переведены по различным причинам.

    Параметры:
        request (Request): Объект запроса.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        TemplateResponse: Отображает данные пациентов в HTML-шаблоне.
    """
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ])
        ))
        .filter(models.PassportData.branch == "Неврологическое отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = False  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Выписанные пациенты Неврологического отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
    
@router.get("/hospitalized_patients_therapeutic_department/")
async def hospitalized_patients_therapeutic_department(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов из терапевтического отделения,
    которые были выписаны или переведены по различным причинам.

    Параметры:
        request (Request): Объект запроса.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        TemplateResponse: Отображает данные пациентов в HTML-шаблоне.
    """
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ])
        ))
        .filter(models.PassportData.branch == "Терапевтическое отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = False  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Выписанные пациенты Терапевтического отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
    
@router.get("/hospitalized_patients_dermatovenerological_department/")
async def hospitalized_patients_dermatovenerological_department(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов из кожно-венерологического отделения,
    которые были выписаны или переведены по различным причинам.

    Параметры:
        request (Request): Объект запроса.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        TemplateResponse: Отображает данные пациентов в HTML-шаблоне.
    """
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ])
        ))
        .filter(models.PassportData.branch == "Кожно-венерологическое отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ])
            )
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = False  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Выписанные пациенты Кожно-венерологического отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
    
@router.get("/hospitalized_patients_surgery_department/")
async def hospitalized_patients_surgery_department(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов из хирургического отделения,
    которые были выписаны или переведены по различным причинам.

    Параметры:
        request (Request): Объект запроса.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        TemplateResponse: Отображает данные пациентов в HTML-шаблоне.
    """
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ])
        ))
        .filter(models.PassportData.branch == "Хирургическое отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ])
            )
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = False  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Выписанные пациенты Хирургического отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
    
#endregion

# region active at branch

'''Пациенты на госпитализации'''
def create_patient_document(patient_data_with_completeness, today, branch):
    # Создаем новый документ Word
    doc = Document()

    # Добавляем заголовок
    title = f"Пациенты находящиеся на госпитализации в {branch} на {today}"
    doc.add_heading(title, level=1)

    # Создаем таблицу
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Устанавливаем шрифт Times New Roman для всего документа
    for paragraph in doc.paragraphs:
        paragraph.style.font.name = 'Times New Roman'
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'

    # Добавляем заголовки столбцов
    headers = ["Номер палаты", "Воинское звание", "ФИО", "Дата рождения", "Дата поступления", "Основа службы"]
    row = table.rows[0].cells
    for i, header in enumerate(headers):
        row[i].text = header

    # Добавляем данные пациентов
    for passport_data, hospital_data, completeness_percentage, color, discharge_date in patient_data_with_completeness:
        # Разбиение полного имени на отдельные части
        parts = passport_data.full_name.split()
        if len(parts) == 3:
            last_name, first_name, patronymic = parts
            full_name = f"{last_name} {first_name[0]}.{patronymic[0]}."
        else:
            full_name = passport_data.full_name  # Оставляем полное имя, если не удалось разбить на три части
        
        # Добавляем строку в таблицу с данными пациента
        row = table.add_row().cells
        row[0].text = ""  # Номер палаты остается пустым
        row[1].text = passport_data.military_rank
        row[2].text = full_name
        row[3].text = passport_data.birthday_date.strftime('%d.%m.%Y') if passport_data.birthday_date else ''
        row[4].text = passport_data.current_time.strftime('%d.%m.%Y %H:%M') if passport_data.current_time else ''
        row[5].text = passport_data.service_basis

    # Сохраняем документ Word
    file_name = f"Пациенты находящиеся на госпитализации {today}.docx"
    doc_path = os.path.join(docs_dir, file_name)
    doc.save(doc_path)

    return doc_path

@router.get("/dermatovenerological_department_active_wards/")
async def dermatovenerological_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Кожно-венерологическое отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .all()
    )

    

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    today = datetime.now().strftime('%Y-%m-%d')
    
    # Создаем документ Word с данными пациентов
    doc_path = create_patient_document(patient_data_with_completeness, today, "Кожно-венерологическом отделении")

    # Отправляем файл пользователю
    response = FileResponse(doc_path, filename=os.path.basename(doc_path))

    return response

@router.get("/surgery_department_active_wards/")
async def dermatovenerological_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Хирургическое отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    today = datetime.now().strftime('%Y-%m-%d')
    
    # Создаем документ Word с данными пациентов
    doc_path = create_patient_document(patient_data_with_completeness, today, "Хирургическом отделении")

    # Отправляем файл пользователю
    response = FileResponse(doc_path, filename=os.path.basename(doc_path))

    return response

@router.get("/neurological_department_active_wards/")
async def dermatovenerological_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Неврологическое отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    today = datetime.now().strftime('%Y-%m-%d')
    
    # Создаем документ Word с данными пациентов
    doc_path = create_patient_document(patient_data_with_completeness, today, "Неврологическом отделении")

    # Отправляем файл пользователю
    response = FileResponse(doc_path, filename=os.path.basename(doc_path))

    return response

@router.get("/therapeutic_department_active_wards/")
async def dermatovenerological_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Терапевтическое отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    today = datetime.now().strftime('%Y-%m-%d')
    
    # Создаем документ Word с данными пациентов
    doc_path = create_patient_document(patient_data_with_completeness, today, "Терапевтическом отделении")

    # Отправляем файл пользователю
    response = FileResponse(doc_path, filename=os.path.basename(doc_path))

    return response

@router.get("/infectious_department_active_wards/")
async def dermatovenerological_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Инфекционное отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    today = datetime.now().strftime('%Y-%m-%d')
    
    # Создаем документ Word с данными пациентов
    doc_path = create_patient_document(patient_data_with_completeness, today, "Инфекционном отделении")

    # Отправляем файл пользователю
    response = FileResponse(doc_path, filename=os.path.basename(doc_path))

    return response



@router.get("/dermatovenerological_department_active/")
async def dermatovenerological_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Кожно-венерологическое отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)


        .all()
    )

     
    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": f"Пациенты, находящиеся на госпитализации в Кожно-венерологическом отделении",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

@router.get("/surgery_department_active/")
async def surgery_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Хирургическое отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)


        .all()
    )

     
    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": f"Пациенты, находящиеся на госпитализации в Хирургическом отделении",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

@router.get("/neurological_department_active/")
async def neurological_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Неврологическое отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)


        .all()
    )

     
    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": f"Пациенты, находящиеся на госпитализации в Неврологическом отделении",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

@router.get("/therapeutic_department_active/")
async def therapeutic_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Терапевтическое отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)


        .all()
    )

     
    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,    
        "title": f"Пациенты, находящиеся на госпитализации в Терапевтическом отделении",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

@router.get("/infectious_diseases_department_active/")
async def therapeutic_department_active(request: Request, date: str = None, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
            or_(
                models.PatientMovement.event_type == "Прием на госпитализацию",
                models.PatientMovement.event_type == "Перевод в другое отделение",
                models.PatientMovement.event_type == "Планируется перевод"
            )
        ))
        .filter(models.PassportData.branch == "Инфекционное отделение")
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)


        .all()
    )

     
    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки пациента, если она есть
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, models.PatientMovement.event_type == "Выписан")
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,    
        "title": f"Пациенты, находящиеся на госпитализации в Инфекционном отделении",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })

#endregion

# region archive

@router.get("/dermatovenerological_department/")
async def hospitalized_patients_infectious_diseases_department(request: Request, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .filter(models.PassportData.branch == "Кожно-венерологическое отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        discharge_date = format_datetime(discharge_date.event_date) if discharge_date else None

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = False  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Архив Кожно-венерологического отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
  
@router.get("/surgery_department/")
async def hospitalized_patients_infectious_diseases_department(request: Request, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .join(models.PatientMovement, and_(
            models.PatientMovement.patient_id == subquery.c.patient_id,
            models.PatientMovement.event_date == subquery.c.latest_event_date,
        ))
        .filter(models.PassportData.branch == "Хирургическое отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        discharge_date = format_datetime(discharge_date.event_date) if discharge_date else None

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Архив Хирургического отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
  
@router.get("/neurological_department/")
async def hospitalized_patients_infectious_diseases_department(request: Request, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .filter(models.PassportData.branch == "Неврологическое отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        discharge_date = format_datetime(discharge_date.event_date) if discharge_date else None

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Архив пациентов Неврологического отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
  
@router.get("/therapeutic_department/")
async def hospitalized_patients_infectious_diseases_department(request: Request, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .filter(models.PassportData.branch == "Терапевтическое отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        discharge_date = format_datetime(discharge_date.event_date) if discharge_date else None

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Архив пациентов Терапевтического отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
    
@router.get("/infectious_diseases_department/")
async def hospitalized_patients_infectious_diseases_department(request: Request, db: Session = Depends(get_db)):
    subquery = db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_event_date')
    ).group_by(models.PatientMovement.patient_id).subquery()

    patients_admitted_on_date = (
        db.query(models.PassportData, models.HospitalData)
        .join(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .join(subquery, models.PassportData.id == subquery.c.patient_id)
        .filter(models.PassportData.branch == "Инфекционное отделение")
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    patient_data_with_completeness = []
    for passport_data, hospital_data in patients_admitted_on_date:
        # Расчет процента заполненности данных пациента
        incomplete_fields, completeness_percentage = check_patient_data_completeness(passport_data)

        # Определение цветового кода в зависимости от процента заполненности
        if completeness_percentage < 45:
            color = 'red'
        elif completeness_percentage < 70:
            color = 'yellow'
        else:
            color = 'green'

        # Получение даты выписки, если пациент выписан
        discharge_date = (
            db.query(models.PatientMovement.event_date)
            .filter(models.PatientMovement.patient_id == passport_data.id, 
            models.PatientMovement.event_type.in_([
                "Выписан",
                "Перевод в другое ВМО",
                "Направлен в санаторий",
                "СОЧ",
                "Летальный исход",
                "Перевод в РКБ",
                "Перевод в другое МО",
                "Выписан за нарушение режима"
            ]))
            .order_by(models.PatientMovement.event_date.desc())
            .first()
        )

        discharge_date = format_datetime(discharge_date.event_date) if discharge_date else None

        # Добавление обработанных данных в список
        patient_data_with_completeness.append((passport_data, hospital_data, completeness_percentage, color, discharge_date))

    # Определение флага для скрытия или отображения столбца с датой выписки
    hide_discharge_date = True  # Измените это значение в зависимости от ваших условий

    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("patient_hospital/patient_table.html", {
        "request": request,
        "passport_data": patient_data_with_completeness,
        "title": "Архив пациентов Инфекционного отделения",
        "selected_date": datetime.now().strftime('%Y-%m-%d'),
        "hide_discharge_date": hide_discharge_date
    })
 
# endregion

# region vision

@router.get("/create_form/")
async def passport_data_html(request: Request, db: Session = Depends(get_db)):
    passport_data = db.query(models.PassportData).all()
    return templates.TemplateResponse("passport_data_form.html", {"request": request, "passport_data": passport_data})

@router.get("/vvk_patients/", response_class=HTMLResponse)
async def vvk_patients(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов вместе с их
    госпитальными записями и данными ВВК.

    Параметры:
        request (Request): Объект запроса.
        db (Session): Зависимость сессии базы данных.

    Возвращает:
        TemplateResponse: Отображает данные пациентов в HTML-шаблоне.
    """
    # Запрос для получения всех пациентов с их госпитальными данными и данными ВВК
    patients_with_vvc_data: List[Tuple[models.PassportData, models.HospitalData, Optional[models.VvkHospital]]] = (
        db.query(models.PassportData, models.HospitalData, models.VvkHospital)
        .outerjoin(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .outerjoin(models.VvkHospital, models.VvkHospital.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )


    # Рендеринг шаблона с предоставленным контекстом
    return templates.TemplateResponse("/patient_hospital/vvk_patients.html", {
        "request": request,
        "passport_data": patients_with_vvc_data,
        "title": "Заключения ВВК",
    })

@router.get("/severity_certificates/", response_class=HTMLResponse)
async def severity_certificates(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов вместе с их
    сертификатами тяжести увечья.
    """
    patients_with_severity_certificates = (
        db.query(models.PassportData, models.HospitalData, models.CertificateOfSeverity)
        .outerjoin(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .outerjoin(models.CertificateOfSeverity, models.CertificateOfSeverity.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    return templates.TemplateResponse("/patient_hospital/severity_certificates.html", {
        "request": request,
        "passport_data": patients_with_severity_certificates,
        "title": "Справки о тяжести увечья",
    })

@router.get("/injury_certificates/", response_class=HTMLResponse)
async def injury_certificates(request: Request, db: Session = Depends(get_db)):
    """
    Асинхронный маршрут для получения и отображения данных пациентов вместе с их
    сертификатами травм.
    """
    patients_with_injury_certificates = (
        db.query(models.PassportData, models.HospitalData, models.CertificateOfInjury)
        .outerjoin(models.HospitalData, models.HospitalData.patient_id == models.PassportData.id)
        .outerjoin(models.CertificateOfInjury, models.CertificateOfInjury.patient_id == models.PassportData.id)
        .order_by(models.PassportData.current_time.desc())
        .all()
    )

    return templates.TemplateResponse("/patient_hospital/injury_certificates.html", {
        "request": request,
        "passport_data": patients_with_injury_certificates,
        "title": "Справки о ранении",
    })

@router.get("/passport_data/{patient_id}", response_class=HTMLResponse)
def read_patient_data(request: Request, patient_id: int, db: Session = Depends(get_db)):
    # Запрос данных о пациенте и его связанных записях в одной транзакции
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if patient is None:
        raise HTTPException(status_code=404, detail="Patient not found")

    # Если пациент найден, получаем связанные данные
    certificateOfSeverity_record = db.query(models.CertificateOfSeverity).filter(models.CertificateOfSeverity.patient_id == patient_id).all()
    certificateOfInjury = db.query(models.CertificateOfInjury).filter(models.CertificateOfInjury.patient_id == patient_id).all()
    vvk_records = db.query(models.VvkHospital).filter(models.VvkHospital.patient_id == patient_id).all()
    hospital_records = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).all()
    movements = db.query(models.PatientMovement).filter(models.PatientMovement.patient_id == patient_id).all()

    therapists = db.query(models.HospitalData.therapist).distinct().all()
    therapists = [therapist[0] for therapist in therapists]  

    therapists_data = db.query(models.HospitalRecord).distinct().all()
    # Возвращаем данные через шаблон
    return templates.TemplateResponse("pacient_data.html", {
        "request": request,
        "passport_data": patient,
        "hospital_records": hospital_records,
        "movements": movements,
        "vvk_records": vvk_records,
        "certificateOfSeverity_record": certificateOfSeverity_record,
        "certificateOfInjury": certificateOfInjury,
        "therapists": therapists,
        "therapists_data": therapists_data
    })


#endregion

# region control

@router.post('/update_diagnosis_record/{record_id}')
def update_diagnosis_record(
    request: Request,
    record_id : int, 
    patient_id: int = Form(...),
    expert_diagnosis: Optional[str] = Form(None),
    final_diagnosis: Optional[str] = Form(None),
    ICD: Optional[str] = Form(None),
    vvk_decision: Optional[str] = Form(None),
    severity_of_injury: Optional[bool] = Form(False),
    anamnesis: Optional[str] = Form(None),
    therapist: Optional[str]= Form(None),
    db: Session = Depends(get_db)):

    record = db.query(models.HospitalData).filter(models.HospitalData.id == record_id).first()
    if record is None:
        raise HTTPException(status_code=404, detail="Patient not found")
    # Обновите запись в базе данных с использованием record_id
    # Например:
    
    record.expert_diagnosis = expert_diagnosis
    record.final_diagnosis = final_diagnosis
    record.ICD = ICD
    record.vvk_decision = vvk_decision
    record.severity_of_injury = severity_of_injury
    record.anamnesis = anamnesis
    record.therapist = therapist

    db.commit()
    return RedirectResponse(url=f"/pacients/passport_data/{patient_id}", status_code=303)

@router.post('/update_hospital_record/{record_id}')
def update_hospital_record(request: Request,
    record_id,
    patient_id: int = Form(...),
    diagnosis_upon_admission: Optional[str] = Form(None),
    vk_urgent_call_date: Optional[str] = Form(None),
    vk_call_up_date: Optional[str] = Form(None),
    district: Optional[str] = Form(None),
    character_of_the_hospital: Optional[str] = Form(None),
    reason_for_departure: Optional[str] = Form(None),
    certificate_of_injury: Optional[bool] = Form(False),
    medical_record: Optional[bool] = Form(False),
    food_certificate: Optional[bool] = Form(False),
    entered_after_participating_in_hostilities: Optional[bool] = Form(False),
    suitability_category: Optional[bool] = Form(False),
    diagnosis_according_to_form_one_hundred: Optional[str] = Form(None),
    db: Session = Depends(get_db)):
    
    record = db.query(models.HospitalData).filter(models.HospitalData.id == record_id).first()
    if record is None:
        raise HTTPException(status_code=404, detail="Patient not found")

    record.diagnosis_upon_admission = diagnosis_upon_admission
    record.vk_urgent_call_date = vk_urgent_call_date
    record.vk_call_up_date = vk_call_up_date
    record.district = district
    record.character_of_the_hospital = character_of_the_hospital
    record.reason_for_departure = reason_for_departure
    record.certificate_of_injury = certificate_of_injury
    record.medical_record = medical_record
    record.food_certificate = food_certificate
    record.entered_after_participating_in_hostilities = entered_after_participating_in_hostilities
    record.suitability_category = suitability_category
    record.diagnosis_according_to_form_one_hundred = diagnosis_according_to_form_one_hundred

    db.commit()
    return RedirectResponse(url=f"/pacients/passport_data/{patient_id}", status_code=303)

@router.post("/update_hospital_data/{record_id}")
async def update_hospital_data(    
    record_id: int,
    patient_id: Optional[str] = Form(...),
    vk_urgent_call_date: Optional[str] = Form(None),
    vk_call_up_date: Optional[str] = Form(None),
    district: Optional[str] = Form(None),
    diagnosis_upon_admission: Optional[str] = Form(None),
    final_diagnosis: Optional[str] = Form(None),
    ICD: Optional[str] = Form(None),
    character_of_the_hospital: Optional[str] = Form(None),
    reason_for_departure: Optional[str] = Form(None),
    vvk_decision: Optional[str] = Form(None),
    certificate_of_injury: Optional[bool] = Form(False),
    food_certificate: Optional[bool] = Form(False),
    medical_record: Optional[bool] = Form(False),
    anamnesis: Optional[str] = Form(None),
    severity_of_injury: Optional[bool] = Form(False),
    entered_after_participating_in_hostilities: Optional[bool] = Form(False),
    db: Session = Depends(get_db)):
    record = db.query(models.HospitalData).filter(models.HospitalData.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")

    record.vk_urgent_call_date = vk_urgent_call_date
    record.vk_call_up_date = vk_call_up_date
    record.district = district
    record.diagnosis_upon_admission = diagnosis_upon_admission
    record.final_diagnosis = final_diagnosis
    record.ICD = ICD
    record.character_of_the_hospital = character_of_the_hospital
    record.reason_for_departure = reason_for_departure
    record.vvk_decision = vvk_decision
    record.certificate_of_injury = certificate_of_injury
    record.food_certificate = food_certificate
    record.medical_record = medical_record
    record.severity_of_injury = severity_of_injury
    record.anamnesis = anamnesis
    record.entered_after_participating_in_hostilities = entered_after_participating_in_hostilities

    db.commit()
    return RedirectResponse(url=f"/pacients/passport_data/{patient_id}", status_code=303)

@router.post("/update_passport_data/{patient_id}")
async def update_passport_data(
    patient_id: int, 
    full_name: Optional[str] = Form(None),
    birthday_date: Optional[str] = Form(None),
    military_rank: Optional[str] = Form(None),
    military_unit: Optional[str] = Form(None),
    service_basis: Optional[str] = Form(None),
    directions: Optional[str] = Form(None),
    branch: Optional[str] = Form(None),
    history_number: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    phone_number: Optional[str] = Form(None),
    personal_data: Optional[str] = Form(None),
    personal_document: Optional[str] = Form(None),
    date_of_illness: Optional[str] = Form(None),
    first_diagnosis: Optional[str] = Form(None),
    after_hostilities: Optional[bool] = Form(False),
    unit_commander: Optional[str] = Form(False),
    db: Session = Depends(get_db)):

    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if patient is None:
        raise HTTPException(status_code=404, detail="Patient not found")

    # Обновляем данные пациента
    patient.full_name = full_name
    patient.birthday_date = birthday_date
    patient.military_rank = military_rank
    patient.military_unit = military_unit
    patient.service_basis = service_basis
    patient.directions = directions
    patient.branch = branch
    patient.history_number = history_number
    patient.address = address
    patient.phone_number = phone_number
    patient.personal_data = personal_data
    patient.personal_document = personal_document
    patient.date_of_illness = date_of_illness
    patient.after_hostilities = after_hostilities
    patient.first_diagnosis = first_diagnosis
    patient.unit_commander = unit_commander

    db.commit()
    
    # Редирект на страницу с формой
    return RedirectResponse(url=f"/pacients/passport_data/{patient_id}", status_code=303)

@router.post("/passport_data/", response_model=schemas.PassportData)
def create_passport_data(
    full_name: str = Form(...),
    birthday_date: Optional[str] = Form(None),
    personal_data: Optional[str] = Form(None),
    military_rank: Optional[str] = Form(None),
    directions: Optional[str] = Form(None),
    date_of_illness: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    military_unit: Optional[str] = Form(None),
    history_number: Optional[str] = Form(None),
    phone_number: Optional[str] = Form(None),
    branch: Optional[str] = Form(None),
    service_basis: Optional[str] = Form(None),
    personal_document: Optional[str] = Form(None),
    first_diagnosis: Optional[str] = Form(None),
    after_hostilities: Optional[bool] = Form(False),
    unit_commander: Optional[str] = Form(None),
    db: Session = Depends(get_db)):
    passport_data_dict = {
        "full_name": full_name,
        "birthday_date": datetime.strptime(birthday_date, "%Y-%m-%d") if birthday_date else None,
        "personal_data": personal_data,
        "military_rank": military_rank,
        "directions": directions,
        "date_of_illness": datetime.strptime(date_of_illness, "%Y-%m-%d") if date_of_illness else None,
        "address": address,
        "military_unit": military_unit,
        "history_number": history_number,
        "phone_number": phone_number,
        "branch": branch,
        "service_basis": service_basis,
        "personal_document" : personal_document,
        "after_hostilities": after_hostilities,
        "first_diagnosis": first_diagnosis,
        'unit_commander': unit_commander,
        'current_time': get_moscow_time()
    }

    db_passport_data = models.PassportData(**passport_data_dict)
    db.add(db_passport_data)
    db.commit()
    db.refresh(db_passport_data)

    return db_passport_data

@router.delete("/passport_data/{passport_data_id}")
def delete_passport_data(passport_data_id: int, db: Session = Depends(get_db)):
    db_passport_data = db.query(models.PassportData).filter(models.PassportData.id == passport_data_id).first()
    if db_passport_data is None:
        raise HTTPException(status_code=404, detail="Passport data not found")
    db.delete(db_passport_data)
    db.commit()
    return {"message": "Passport data deleted successfully"}

@router.delete("/delete_all_patient_data/{patient_id}", dependencies=[Depends(check_ip)])
async def delete_all_patient_data(patient_id: int, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    
    db.delete(patient)
    db.commit()
    return {"message": "All patient data deleted successfully"}

@router.post("/create_patient_vvk", dependencies=[Depends(check_ip_vvk)])
async def create_patient_movement(
    patient_id: int = Form(...),
    vvk_number: str = Form(...),
    conclusion: str = Form(...),
    vvk_date: datetime = Form(...),
    date_of_dispatch: datetime = Form(None),
    date_of_approval: datetime = Form(None),
    db: Session = Depends(get_db)):

    # Создание новой записи о перемещении
    new_movement = models.VvkHospital(
        patient_id=patient_id,
        vvk_number=vvk_number,
        conclusion=conclusion,
        vvk_date=vvk_date,
        date_of_dispatch=date_of_dispatch,
        date_of_approval=date_of_approval
    )
    db.add(new_movement)

    # Попытка зафиксировать изменения в базе данных
    try:
        db.commit()
        return JSONResponse(status_code=200, content={"message": "Информация о заключении ВВК успешно добавлена."})
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/create_patient_injury", dependencies=[Depends(check_ip_vvk)])
async def create_patient_movement(
    patient_id: int = Form(...),
    load_date: Optional[datetime] = Form(None),
    injury_number: str = Form(...),
    certificate_injury_date: datetime = Form(...),
    db: Session = Depends(get_db)):

    # Создание новой записи о перемещении
    new_movement = models.CertificateOfInjury(
        patient_id=patient_id,
        load_date=load_date,
        injury_number=injury_number,
        certificate_injury_date=certificate_injury_date
    )
    db.add(new_movement)

    # Попытка зафиксировать изменения в базе данных
    try:
        db.commit()
        return JSONResponse(status_code=200, content={"message": "Справка о ранении успешно добавлена."})
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/create_patient_severity", dependencies=[Depends(check_ip_vvk)])
async def create_patient_movement(
    patient_id: int = Form(...),
    approval_date: Optional[datetime] = Form(None),
    approval_number: Optional[str] = Form(None),
    severity_number: str = Form(...),
    severity_date: datetime = Form(...),
    db: Session = Depends(get_db)):

    # Создание новой записи о перемещении
    new_movement = models.CertificateOfSeverity(
        patient_id=patient_id,
        approval_date=approval_date,
        approval_number=approval_number,
        severity_number=severity_number,
        severity_date=severity_date
    )
    db.add(new_movement)

    # Попытка зафиксировать изменения в базе данных
    try:
        db.commit()
        return JSONResponse(status_code=201, content={"message": "Справка о тяжести увечья успешно добавлена."})
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

import pytz

@router.post("/create_patient_movement")
async def create_patient_movement(
    patient_id: int = Form(...),
    department: str = Form(...),
    event_type: str = Form(...),
    event_date: datetime = Form(...),
    destination_department: Optional[str] = Form(None),
    note: Optional[str] = Form(None),
    db: Session = Depends(get_db)):

    # Приведение event_date к UTC
    if event_date.tzinfo is None:
        event_date = pytz.utc.localize(event_date) - timedelta(hours=3)
    else:
        event_date = event_date.astimezone(pytz.utc) - timedelta(hours=3)

    # Найти последнее перемещение для данного пациента
    last_movement = db.query(models.PatientMovement).filter(
        models.PatientMovement.patient_id == patient_id
    ).order_by(models.PatientMovement.event_date.desc()).first()

    if last_movement:
        # Приведение event_date последнего перемещения к UTC
        if last_movement.event_date.tzinfo is None:
            last_movement.event_date = pytz.utc.localize(last_movement.event_date)
        else:
            last_movement.event_date = last_movement.event_date.astimezone(pytz.utc)

        # Если новое перемещение происходит до последнего, скорректировать его время
        if event_date <= last_movement.event_date:
            event_date = last_movement.event_date + timedelta(minutes=1) 

    # Создать новую запись перемещения
    new_movement = models.PatientMovement(
        patient_id=patient_id,
        department=department,
        event_type=event_type,
        event_date=event_date,
        destination_department=destination_department,
        note=note
    )
    db.add(new_movement)

    # Если тип события "Перевод в другое отделение", обновить данные в PassportData
    if event_type == "Перевод в другое отделение" and destination_department:
        passport_data = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
        if passport_data:
            passport_data.branch = destination_department
        else:
            db.rollback()
            raise HTTPException(status_code=404, detail="Patient not found")

    # Попытаться сохранить изменения в базе данных
    try:
        db.commit()
        return RedirectResponse(url=f"/pacients/passport_data/{patient_id}", status_code=303)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    
@router.delete("/delete_patient_movement/{patient_id}/{movement_id}", dependencies=[Depends(check_ip)])
async def delete_patient_movement(request: Request, patient_id: int, movement_id: int, db: Session = Depends(get_db)):
    movement = db.query(models.PatientMovement).filter(models.PatientMovement.id == movement_id).first()
    if not movement:
        raise HTTPException(status_code=404, detail="Patient movement not found")
    db.delete(movement)
    db.commit()
    return RedirectResponse(url=f"/pacients/passport_data/{patient_id}", status_code=303)

@router.delete("/delete_patient_vvk/{patient_id}/{vvk_id}", dependencies=[Depends(check_ip_vvk)])
async def delete_patient_vvk(request: Request, patient_id: int, vvk_id: int, db: Session = Depends(get_db)):
    movement = db.query(models.VvkHospital).filter(models.VvkHospital.id == vvk_id).first()
    if not movement:
        raise HTTPException(status_code=404, detail="Patient movement not found")
    db.delete(movement)
    db.commit()
    return JSONResponse(status_code=200, content={"message": "Заключение ВВК успешно удалено."})


@router.delete("/delete_patient_injury/{patient_id}/{injury_id}", dependencies=[Depends(check_ip_vvk)])
async def delete_patient_injury(request: Request, patient_id: int, injury_id: int, db: Session = Depends(get_db)):
    movement = db.query(models.CertificateOfInjury).filter(models.CertificateOfInjury.id == injury_id).first()
    if not movement:
        raise HTTPException(status_code=404, detail="Patient movement not found")
    db.delete(movement)
    db.commit()
    return JSONResponse(status_code=200, content={"message": "Удаление справки о ранении успешно выполнено."})

@router.delete("/delete_patient_severity/{patient_id}/{severity_id}", dependencies=[Depends(check_ip_vvk)])
async def delete_patient_severity(request: Request, patient_id: int, severity_id: int, db: Session = Depends(get_db)):
    movement = db.query(models.CertificateOfSeverity).filter(models.CertificateOfSeverity.id == severity_id).first()
    if not movement:
        raise HTTPException(status_code=404, detail="Patient movement not found")
    db.delete(movement)
    db.commit()
    return JSONResponse(status_code=200, content={"message": "Удаление справки о тяжести увечья успешно выполнено."})


@router.put("/update_certificate_of_severity/{severity_id}", dependencies=[Depends(check_ip)])
async def update_certificate_of_severity(severity_id: int,
                                         severity_number: str = Form(...),
                                         severity_date: datetime = Form(...),
                                         approval_date: Optional[datetime] = Form(None),
                                         approval_number: Optional[str] = Form(None),
                                         patient_id: int = Form(...),
                                         db: Session = Depends(get_db)):
    severity = db.query(models.CertificateOfSeverity).filter(models.CertificateOfSeverity.id == severity_id).first()
    if not severity:
        raise HTTPException(status_code=404, detail="Certificate of severity not found")

    severity.severity_number = severity_number
    severity.severity_date = severity_date
    severity.approval_date = approval_date
    severity.approval_number = approval_number
    severity.patient_id = patient_id
    db.commit()
    return {"message": "Certificate of severity updated successfully"}

@router.put("/update_certificate_of_injury/{injury_id}", dependencies=[Depends(check_ip_vvk)])
async def update_certificate_of_injury(injury_id: int,
                                       injury_number: str = Form(...),
                                       certificate_injury_date: datetime = Form(...),
                                       load_date: datetime = Form(...),
                                       patient_id: int = Form(...),
                                       db: Session = Depends(get_db)):
    injury = db.query(models.CertificateOfInjury).filter(models.CertificateOfInjury.id == injury_id).first()
    if not injury:
        raise HTTPException(status_code=404, detail="Certificate of injury not found")

    injury.injury_number = injury_number
    injury.certificate_injury_date = certificate_injury_date
    injury.load_date = load_date
    injury.patient_id = patient_id
    db.commit()
    return {"message": "Certificate of injury updated successfully"}


@router.put("/patient_movements/{movement_id}", response_model=schemas.PatientMovementCreate)
async def update_patient_movement(movement_id: int, movement_data: schemas.PatientMovementUpdate, db: Session = Depends(get_db)):
    movement = db.query(models.PatientMovement).filter(models.PatientMovement.id == movement_id).first()
    if not movement:
        raise HTTPException(status_code=404, detail="Patient movement not found")
    for var, value in movement_data.dict(exclude_unset=True).items():
        setattr(movement, var, value)
    db.commit()
    return movement

@router.put("/update_patient_movement/{movement_id}")
async def update_patient_movement(movement_id: int, department: str = Form(...),
                                  event_type: str = Form(...), event_date: datetime = Form(...),
                                  destination_department: Optional[str] = Form(None),
                                  db: Session = Depends(SessionLocal)):
    movement = db.query(models.PatientMovement).filter(models.PatientMovement.id == movement_id).first()
    if not movement:
        raise HTTPException(status_code=404, detail="Movement not found")

    movement.department = department
    movement.event_type = event_type
    movement.event_date = event_date
    movement.destination_department = destination_department
    db.commit()
    return {"message": "Patient movement updated successfully"}

#endregion