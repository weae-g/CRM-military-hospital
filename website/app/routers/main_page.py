from fastapi import APIRouter, HTTPException, Depends, Request, BackgroundTasks
from sqlalchemy.orm import Session
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from typing import List
from datetime import datetime, timedelta, time
from sqlalchemy import case, func, and_, or_, extract
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from typing import Dict
from ._utils.db_utils import get_db
import time
import os
from datetime import date
import logging
from docx.shared import Inches
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from celery import Celery
import traceback
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi.responses import StreamingResponse
import urllib.parse
from ..model.database.database import SessionLocal
from ..model.pacient_hospital import models, schemas
from ._utils.check_utils import check_completeness_by_branch, check_patient_data_completeness
from ._utils.documents_utils import current_dir, docs_dir, img_dir, js_dir, css_dir, templates
from ._utils.сases_utils import get_ablt_case, get_accs_case, get_gent_case, get_loct_case, get_dative_case
from ._utils.documents_utils import (
    create_header_with_doctor_and_shape,
    fill_template,
    set_paragraph_formatting,
    add_custom_run,
    add_custom_run_yellow,
    set_cell_border,
    set_cell_vertical_alignment,
    center_text_in_line,
    add_full_underlined_text,
    set_font_to_times_new_roman,
    fill_template_v_times,
    format_name,
    delete_file_with_delay,
    set_cell_border_v2,
    set_table_borders,
    set_font_and_spacing
)
from docx.shared import Cm
from io import BytesIO
import tempfile
import calendar
import locale
from typing import List, Tuple, Optional, Dict, Union
from natasha import NamesExtractor



logger = logging.getLogger(__name__)

router = APIRouter()






#region 
branches_communications = {
        'Кожно-венерологическое отделение' : "ВрИО ГП ВС Кутдусова А.М.", 
        'Неврологическое отделение': "Начальник неврологического отделения ГП ВС Токарева Е.В.",
        'Терапевтическое отделение': 'Начальник терапевтического отделения майор м/с Косов А.А.',
        'Хирургическое отделение': 'Начальник хирургического отделения капитан м/с Неклюдов Д.С.',
        'Инфекционное отделение': 'Начальник инфекционного отделения майор м/с Назимов И.Н.'
    }

branches_communications = {
        'Кожно-венерологическое отделение' : "ГП ВС                                            Кутдусова А.М.", 
        'Неврологическое отделение': "        ГП ВС                                            Токарева Е.В.",
        'Терапевтическое отделение': '        майор м/с     	     			               Косов А.А.',
        'Хирургическое отделение': '          майор м/с                                      Неклюдов Д.С.',
        'Инфекционное отделение': '           майор м/с                                        Назимов И.Н.'
    }

branches_communications_name = {
    'Кожно-венерологическое отделение' : "Кутдусова А.М.", 
    'Неврологическое отделение': "Токарева Е.В.",
    'Терапевтическое отделение': 'Косов А.А.',
    'Хирургическое отделение': 'Неклюдов Д.С.',
    'Инфекционное отделение': 'Назимов И.Н.'
}

branches_communications_rank = {
    'Кожно-венерологическое отделение' : "ГП ВС", 
    'Неврологическое отделение': "ГП ВС",
    'Терапевтическое отделение': 'майор',
    'Хирургическое отделение': 'майор',
    'Инфекционное отделение': 'майор'
}

branches_communications_add = {
    'Кожно-венерологическое отделение' : "ВрИО начальника", 
    'Неврологическое отделение': "Начальник",
    'Терапевтическое отделение': 'Начальник',
    'Хирургическое отделение': 'Начальник',
    'Инфекционное отделение': 'Начальник'
}

month_names = {
    1: "«января»",
    2: "«февраля»",
    3: "«марта»",
    4: "«апреля»",
    5: "«мая»",
    6: "«июня»",
    7: "«июля»",
    8: "«августа»",
    9: "«сентября»",
    10: "«октября»",
    11: "«ноября»",
    12: "«декабря»"
}

additional = {
        "Кожно-венерологическое отделение": ' (на 20 коек)',
        'Инфекционное отделение': ' (на 30 коек)',
        'Неврологическое отделение': ' (на 20 коек)',
        'Терапевтическое отделение': ' (на 35 коек)',
        'Хирургическое отделение': ' (на 40 коек)'
    }

branch_mapping = {
        "Терапевтическое отделение": "терапевтического отделения",
        "Хирургическое отделение": "хирургического отделения",
        "Инфекционное отделение": "инфекционного отделения",
        "Неврологическое отделение": "неврологического отделения",
        "Кожно-венерологическое отделение": "кожно-венерологического отделения"
    }
#endregion

#region Документы


'''Первичный осмотр'''
def create_medical_report(patient_id: int, db: Session):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    doc = Document()
    section = doc.sections[0]  # Получаем первую секцию документа
    section.left_margin = Inches(0.5)  # Установка левого отступа
    section.right_margin = Inches(0.5)  # Установка правого отступа
    section.top_margin = Inches(0.5)  # Установка верхнего отступа
    section.bottom_margin = Inches(0.5)  # Установка нижнего отступа

    # Условие для проверки наличия сертификата о травме
    if hospital_data.certificate_of_injury:
        create_header_with_doctor_and_shape(doc, f"{hospital_data.therapist}","Справка о ранении для выплат ЕДВ в соответсвии с Указом Президента РФ от 05.03.2022г. №98 оформлена", 2)
    else:
        create_header_with_doctor_and_shape(doc, f"{hospital_data.therapist}", "Справка о ранении для выплат ЕДВ в соответсвии с Указом Президента РФ от 05.03.2022г. №98 не оформлена", 2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_custom_run(title, 'Протокол первичного осмотра врачом при приеме больного', bold=True)
    set_paragraph_formatting(title, space_after=Pt(12), space_before=Pt(12), line_spacing=1.5)

    intro = doc.add_paragraph()
    add_custom_run(intro, f"Дата: {patient.current_time.strftime('%d.%m.%Y')} \nВремя: {datetime.now().strftime('%H:%M:%S')}", bold=True)
    set_paragraph_formatting(intro, space_after=Pt(12))

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_custom_run(intro, "Воинское звание: ", bold=False)
    add_custom_run(intro, f"{patient.military_rank}", bold=True)
    add_custom_run(intro, ", Контингент: ", bold=False)
    add_custom_run(intro, f"{patient.service_basis}", bold=True)

    patient_name = doc.add_paragraph()
    patient_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_custom_run(patient_name, f"{patient.full_name}, ", bold=True)
    add_custom_run(patient_name, "Год рождения: ", bold=False)
    add_custom_run(patient_name, f"{patient.birthday_date.strftime('%d.%m.%Y')}\n", bold=True)
    set_paragraph_formatting(patient_name, space_after=Pt(12))

    complaints = doc.add_paragraph()
    add_custom_run(complaints, "ЖАЛОБЫ:", bold=True)
    set_paragraph_formatting(complaints, space_after=Pt(12))

    anamnesis = doc.add_paragraph()
    add_custom_run(anamnesis, "АНАМНЕЗ ЗАБОЛЕВАНИЯ:", bold=True)    
    add_custom_run(anamnesis, f"{hospital_data.anamnesis}")
    set_paragraph_formatting(anamnesis, space_after=Pt(12))

    objective = doc.add_paragraph()
    objective.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_custom_run(objective, "ОБЪЕКТИВНО: ", bold=True)    
    add_custom_run(objective, "Общее состояние пациента:", bold=False, underline=True)
    add_custom_run(objective, " удовлетворительное, сознание ясное. ", bold=False)

    add_custom_run(objective, "Физическое развитие:", bold=False, underline=True)
    add_custom_run(objective, " соответствует полу и возрасту. ", bold=False)

    add_custom_run(objective, "Телосложение:", bold=False, underline=True)
    add_custom_run(objective, " нормостеническое. ", bold=False)

    add_custom_run(objective, "Питание:", bold=False, underline=True)
    add_custom_run(objective, " повышенное, рост 178 см, масса тела 94, ИМТ 29,6 кг/м2. ", bold=False)

    add_custom_run(objective, "Видимые слизистые:", bold=False, underline=True)
    add_custom_run(objective, " физиологической окраски, влажные, чистые. ", bold=False)

    add_custom_run(objective, "Зев: ", bold=False, underline=True)
    add_custom_run(objective, " негиперемирован, миндалины не увеличены, налета нет. ", bold=False)

    add_custom_run(objective, "Лимфатические узлы:", bold=False, underline=True)
    add_custom_run(objective, " не увеличены, безболезненные. ", bold=False)

    add_custom_run(objective, "Щитовидная железа:", bold=False, underline=True)
    add_custom_run(objective, " не увеличена. ", bold=False)

    add_custom_run(objective, "Костно-мышечная система:", bold=False, underline=True)
    add_custom_run(objective, " без патологических изменений. ", bold=False)

    add_custom_run(objective, "Связочно-суставной аппарат:", bold=False, underline=True)
    add_custom_run(objective, " без патологии. ", bold=False)

    add_custom_run(objective, "Органы кровообращения:", bold=False, underline=True)
    add_custom_run(objective, " Область сердца без видимых изменений, границы относительной и абсолютной сердечной тупости в норме. Тоны сердца: звучные, ритм правильный, шумов нет. ЧСС 80 ударов в 1 минуту, ритмичный, удовлетворительного наполнения и напряжения. АД 140/100 мм. рт. ст. ", bold=False)

    add_custom_run(objective, "Органы дыхания:", bold=False, underline=True)
    add_custom_run(objective, " Носовое дыхание свободное. Грудная клетка правильной формы, симметрична. ЧДД 18 в 1 минуту. Перкуторно: ясный легочный звук по всем перкуторным точкам. Аускультативно: дыхание везикулярное по всем полям, хрипов нет. ", bold=False)

    add_custom_run(objective, "Органы пищеварения:", bold=False, underline=True)
    add_custom_run(objective, " Язык влажный, чистый. Живот при пальпации: мягкий, болезненный. Печень: не увеличена, находится под краем реберной дуги справа. Селезенка: не пальпируется. Стул: кашицеобразный, без патологических примесей, регулярный. ", bold=False)

    add_custom_run(objective, "Органы мочеполовой системы:", bold=False, underline=True)
    add_custom_run(objective, " Область почек без патологии, поколачивание по поясничной области безболезненное с обеих сторон. Мочеиспускание свободное, регулярное. ", bold=False)

    set_paragraph_formatting(objective, space_after=Pt(12))

    conclusion = doc.add_paragraph()
    add_custom_run(conclusion, f"ДИАГНОЗ ПРИ ПОСТУПЛЕНИИ: ", bold=True)
    add_custom_run(conclusion, f"{hospital_data.diagnosis_upon_admission}\n")
    add_custom_run(conclusion, f"ПРЕДВАРИТЕЛЬНЫЙ ДИАГНОЗ: ", bold=True)
    add_custom_run(conclusion, f"{hospital_data.final_diagnosis}\n")
    add_custom_run(conclusion, f"КОД ПО МКБ-10: ", bold=True)
    add_custom_run(conclusion, f"{hospital_data.ICD}\n")
    set_paragraph_formatting(conclusion, space_after=Pt(12))

    plan = doc.add_paragraph()
    add_custom_run(plan, f"ПЛАН ОБСЛЕДОВАНИЯ И ЛЕЧЕНИЯ:\n", bold=True)
    add_custom_run(plan, "- Режим общий\n- Диета 15\n- Флюорография\n- ЭКГ\n- ОАК, ОАМ\n")
    set_paragraph_formatting(plan, space_after=Pt(12))

   
    doctor_name_paragraph = doc.add_paragraph()
    doctor_name_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT 

    branch_name = branches_communications[patient.branch]

    #dsadasasfds желудок не голова - переварит все 

    add_custom_run(doctor_name_paragraph, f"{branch_name}", bold=True)

    # Установка форматирования абзаца, если необходимо
    set_paragraph_formatting(doctor_name_paragraph, space_after=Pt(12))

    def format_name(full_name):
        # Разбиваем полное имя по пробелам
        parts = full_name.split()
        if len(parts) == 3:
            # Формируем строку в формате "Фамилия И.О."
            formatted_name = f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
        elif len(parts) == 2:
            # Если у пациента только имя и фамилия
            formatted_name = f"{parts[0]} {parts[1][0]}."
        else:
            # Если что-то нестандартное, вернуть как есть
            formatted_name = full_name
        return formatted_name

    doc_filename = f'Певичный осмотр {format_name(patient.full_name)} {datetime.now().date()}.docx'
    doc_path = os.path.join(docs_dir, doc_filename)
    doc.save(doc_path)

    return doc_path

'''Дневники'''
def create_patient_diary(patient_data: models.PassportData, db: Session):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Pt(20)
    section.bottom_margin = Pt(20)
    section.left_margin = Pt(20)
    section.right_margin = Pt(20)

    movements = db.query(models.PatientMovement).filter_by(patient_id=patient_data.id).all()
    hospital_data = db.query(models.HospitalData).filter_by(patient_id=patient_data.id).first()

    admission_movement = next((m for m in movements if m.event_type == 'Прием на госпитализацию'), None)
    discharge_movement = next((m for m in movements if m.event_type in ['Выписан', 'Перевод в РКБ', 'Направлен в санаторий']), None)

    if not admission_movement:
        raise ValueError("Отсутствуют данные о госпитализации")
    if not discharge_movement:
        raise ValueError("Отсутствуют данные о выписке")

    admission_date = admission_movement.event_date
    discharge_date = discharge_movement.event_date
    diary_content = """Общее состояние удовлетворительное. Сознание ясное. Положение активное. Передвигается на костылях. Жалобы на боли во всех областях позвоночника, общая слабость, тремор нижних конечностей, туловища, ранее тремор головы, нарушение походки. Кожный покров физиологической окраски, чистый. Видимые слизистые полости рта розовые, чистые. Пульс ритмичный, удовлетворительного наполнения и напряжения. Тоны сердца ясные, шумов нет. ЧСС 72 в 1 минуту. АД 130 и 80 мм рт. ст. Аускультативно в легких дыхание везикулярное, проводится над всеми полями, хрипов нет. ЧДД –17 в 1 минуту Язык влажный, чистый. Живот мягкий, при пальпации безболезненный во всех областях. Поколачивание по поясничной области безболезненное с обеих сторон. Физиологические оправления не нарушены."""

    

    current_date = admission_date
    while current_date <= discharge_date:
        if current_date.weekday() == 0 or (current_date == admission_date and current_date.weekday() != 0) or (current_date - admission_date).days % 2 == 0:
            doc.add_paragraph(current_date.strftime('%d.%m.%Y')).bold = True

            chief_signature = doc.add_paragraph()
            chief_signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chief_signature_run = chief_signature.add_run("Осмотр начальником медицинской части" if current_date.weekday() == 0 else "Осмотр начальником отделения")
            chief_signature_run.font.size = Pt(11)
            chief_signature_run.bold = True

            diary_paragraph = doc.add_paragraph(diary_content)
            diary_paragraph.style.font.name = 'Times New Roman'
            diary_paragraph.style.font.size = Pt(11)

            for department in set(m.department for m in movements):
                department_sign = doc.add_paragraph()
                department_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                department_sign_run = department_sign.add_run(f"Начальник {branch_mapping.get(department, department)}, м-р м/с\nЛечащий врач ГП\t\t{hospital_data.therapist}")
                department_sign_run.font.size = Pt(11)
                department_sign_run.bold = True

        current_date += timedelta(days=1)

    file_name = f'Дневники {patient_data.full_name.split()[0]} {datetime.now().strftime("%d.%m.%Y")}.docx'
    doc_path = os.path.join(docs_dir, file_name)
    doc.save(doc_path)
    return doc_path

'''Ежедневный доклад'''
def create_document(departments_data, long_stay_patients, service_basis_data, directions_data):
    # Create a new document
    doc = Document()

    # Set the page orientation to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Set margins
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Add header
    header = section.header
    header_paragraph = header.paragraphs[0]

    header_paragraph.text = f"ЕЖЕДНЕВНЫЙ ДОКЛАД ФГКУ ВОЕННЫЙ ГОСПИТАЛЬ (НА 150 КОЕК, Г. Казань) «354 ВКГ» МО РОССИИ\nна 8:00 {datetime.now().strftime('%d.%m.%Y')} г."
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.style.font.size = Pt(12)

    # Add heading for the first table
    doc.add_heading('Статистика по отделениям', level=1)

    # Add table for bed statistics
    table_beds = doc.add_table(rows=1, cols=13)
    table_beds.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table_beds.rows[0].cells
    headers = [
        'Отделение', 'Развернуто коек', 'Состоит человек', 'Поступило', 
        'Выписано', 'Переведено в другое отделение', 'Из другого отделения', 
        'Переведено в другое МО', 'Состоит', 'Фактическое число', 
        'Процент загрузки', 'Свободных коек', 'С хирургической патологией'
    ]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    for dept in departments_data:
        row_cells = table_beds.add_row().cells
        row_data = [
            dept["department"], dept["beds_deployed"], dept["active_patients"], dept["admitted"], 
            dept["discharged"], dept["transferred_out"], dept["transferred_in"], dept["transferred_to_other_MO"], 
            dept["current_patients"], dept["factual_number"], f"{dept['load_percentage']:.2f}%", dept["free_beds"], 
            dept["surgical_pathology"]
        ]
        for i, data in enumerate(row_data):
            row_cells[i].text = str(data)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    # Set table borders to black
    table_beds.style = 'Table Grid'

        # Add second heading for long stay patients
    doc.add_heading('Пациенты, находящиеся на госпитализации более 30 суток', level=1)

    # Add table for long stay patients
    table_long_stay = doc.add_table(rows=1, cols=5)
    table_long_stay.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table_long_stay.rows[0].cells
    headers = ['Порядковый номер', 'Часть', 'Первичный диагноз', 'Причины длительного пребывания', 'Ориентировачная дата выписки']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    for idx, patient in enumerate(long_stay_patients, start=1):
        row_cells = table_long_stay.add_row().cells
        row_data = [idx, patient.military_unit, patient.diagnosis_upon_admission, "", ""]  # Filling empty columns
        for i, data in enumerate(row_data):
            row_cells[i].text = str(data)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    # Set table borders to black
    table_long_stay.style = 'Table Grid'


    # Add second heading for RKB treatment statistics
    doc.add_heading('Пациенты, находящиеся на лечении в РКБ', level=1)

    # Add table for RKB statistics
    table_rkb = doc.add_table(rows=1, cols=3)
    table_rkb.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table_rkb.rows[0].cells
    headers = ['Отделение', 'Запланировано коек', 'Состоит']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    rkb_departments = [
        {"department": "Травматология и ортопедия", "planned_beds": 160, "current_patients": ""},
        {"department": "Хирургия", "planned_beds": 32, "current_patients": ""},
        {"department": "Абдоминальная хирургия", "planned_beds": 70, "current_patients": ""},
        {"department": "Нейрохирургия", "planned_beds": 20, "current_patients": ""},
        {"department": "Реанимация", "planned_beds": 18, "current_patients": ""}
    ]
    for dept in rkb_departments:
        row_cells = table_rkb.add_row().cells
        row_data = [dept["department"], dept["planned_beds"], dept["current_patients"]]
        for i, data in enumerate(row_data):
            row_cells[i].text = str(data)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    # Set table borders to black
    table_rkb.style = 'Table Grid'

    # Add third heading for service basis statistics
    doc.add_heading('Контингент по основаниям службы', level=1)

    # Add table for service basis statistics
    table_service_basis = doc.add_table(rows=1, cols=2)
    table_service_basis.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table_service_basis.rows[0].cells
    headers = ['Основание службы', 'Количество']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    for item in service_basis_data:
        row_cells = table_service_basis.add_row().cells
        row_data = [item.service_basis, item.count]
        for i, data in enumerate(row_data):
            row_cells[i].text = str(data)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    # Set table borders to black
    table_service_basis.style = 'Table Grid'

    # Add fourth heading for directions statistics
    doc.add_heading('Контингент по направлениям', level=1)

    # Add table for directions statistics
    table_directions = doc.add_table(rows=1, cols=2)
    table_directions.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table_directions.rows[0].cells
    headers = ['Направление', 'Количество']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    for item in directions_data:
        row_cells = table_directions.add_row().cells
        row_data = [item.directions, item.count]
        for i, data in enumerate(row_data):
            row_cells[i].text = str(data)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    # Set table borders to black
    table_directions.style = 'Table Grid'

    return doc

'''Снятие с довольствия'''
def create_discharge_report(db: Session) -> str:
    today = datetime.utcnow().date()
    
    # Query to get discharges and hospitalization dates
    discharges = db.query(models.PatientMovement).join(models.PassportData).filter(
        models.PatientMovement.event_date >= today,
        models.PatientMovement.event_type.in_([
            "Перевод в другое МО", "Выписан", "Выписан за нарушение режима",
            "Перевод в другое ВМО", "Перевод в РКБ", "Направлен в санаторий",
            "СОЧ", "Направлен в санаторий", "Летальный исход"
        ])
    ).all()

    # Create the document
    doc = Document()
    heading = doc.add_heading(f"Снятие с колового довольствования на {today.strftime('%d.%m.%Y')}", level=1)

    # Set the font for the heading
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Add table with specific columns
    table = doc.add_table(rows=1, cols=12)  # Added column for 'Койко-день'
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['№', 'Дата', 'Время', 'Ф.И.О', 'В/ЗВ', 'В/Ч', 'Отделение', 'По факту', 'Диета', 'Продаттестат', 'Иное', 'Койко-день']
    
    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        set_cell_border_v2(hdr_cells[i])

    # Adjust column widths
    for col in table.columns:
        for cell in col.cells:
            cell.width = Pt(40)  # Set a narrow width, adjust if necessary

    for i, discharge in enumerate(discharges, start=1):
        row_cells = table.add_row().cells
        admission_date = db.query(models.PatientMovement).filter(
            models.PatientMovement.patient_id == discharge.patient_id,
            models.PatientMovement.event_type == "Прием на госпитализацию"
        ).order_by(models.PatientMovement.event_date).first()
        
        admission_date = admission_date.event_date if admission_date else discharge.event_date

        # Calculate 'Койко-день'
        days_in_hospital = (discharge.event_date - admission_date).days

        row_data = [
            str(i),
            discharge.event_date.strftime('%d.%m.%Y'),
            discharge.event_date.strftime('%H:%M'),
            discharge.patient.full_name,
            discharge.patient.military_rank or '',
            discharge.patient.military_unit or '',
            discharge.department or '',
            discharge.destination_department or '',
            '',  # Диета, заполните по необходимости
            '',  # Продаттестат, заполните по необходимости
            discharge.event_type,  # Например, тип события
            str(days_in_hospital)  # Койко-день
        ]
        for j, data in enumerate(row_data):
            row_cells[j].text = data
            run = row_cells[j].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            set_cell_border_v2(row_cells[j])

    filename = f"Снятие с колового довольствования {today.strftime('%d_%m_%Y')}.docx"
    filepath = os.path.join(docs_dir, filename)
    doc.save(filepath)

    return filepath

'''Ординаторское требование'''
def create_discharge_report_v2(db: Session) -> str:
    today = datetime.utcnow().date()
    
    discharges = db.query(models.PatientMovement).join(models.PassportData).filter(
        models.PatientMovement.event_date >= today,
        models.PatientMovement.event_type.in_(["Прием на госпитализацию"])
    ).all()
    
    doc = Document()
    heading = doc.add_heading(f"Ординаторское требование на {today.strftime('%d.%m.%Y')}", level=1)

    # Установка шрифта Times New Roman для заголовка
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    table = doc.add_table(rows=1, cols=11)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    headers = ['№', 'Дата', 'Время', 'Ф.И.О', 'В/ЗВ', 'В/Ч', 'Отделение', 'По факту', 'Диета', 'Продаттестат', 'Иное']
    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        set_cell_border_v2(hdr_cells[i])

    for i, discharge in enumerate(discharges, start=1):
        row_cells = table.add_row().cells
        row_data = [
            str(i),
            discharge.event_date.strftime('%d.%m.%Y'),
            discharge.event_date.strftime('%H:%M'),
            discharge.patient.full_name,
            discharge.patient.military_rank or '',
            discharge.patient.military_unit or '',
            discharge.department or '',
            discharge.destination_department or '',
            '',  # Диета, заполните по необходимости
            '',  # Продаттестат, заполните по необходимости
            discharge.event_type  # Например, тип события
        ]
        for j, data in enumerate(row_data):
            row_cells[j].text = data
            run = row_cells[j].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            set_cell_border_v2(row_cells[j])

    filename = f"Ординаторское требование на {today.strftime('%d_%m_%Y')}.docx"
    filepath = os.path.join(docs_dir, filename)
    doc.save(filepath)

    return filepath


@router.get("/vvk-report/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    template_path = os.path.join(docs_dir, "templates\\_Все для ВВК.docx")

    service_base = None
    if patient.service_basis == 'Военнослужащий по контракту':
        service_base = "по контракту"
    elif patient.service_basis == "Военнослужащий, призванный мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Военнослужащий, призванный по мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Курсант 1 курса (призывник)" or patient.service_basis == "Военнослужащий по призыву":
        service_base = "по призыву"
    else:
        service_base = "по контракту"

    
   
    patient_current_month = datetime.now().month
    month_name = month_names[patient_current_month]

    gent_branch = get_gent_case(patient.branch).lower()

    context = {
        'full_name': patient.full_name.upper(),
        'full_name_p2': patient.full_name,
        'today': datetime.now().strftime('%d {} %Y').format(month_name),
        'full_name_p1': get_accs_case(patient.full_name) if get_accs_case(patient.full_name).count == patient.full_name.count else patient.full_name,
        'branch_p': f'{branches_communications_add[patient.branch]} {gent_branch}' if gent_branch else patient.branch.lower(),
        'branch_rank': branches_communications_rank[patient.branch],
        'name_p1': branches_communications_name[patient.branch],
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_basis': service_base.lower() if service_base.lower() else service_base,
        'birthday': patient.birthday_date.strftime('%d.%m.%Y'),
        'finally_diagnos': hospital_data.expert_diagnosis,
        'vk_urgent_call_date': f'Служба по призыву: {hospital_data.vk_call_up_date}' if hospital_data.vk_call_up_date else 'Служба по призыву: не служил',
        'vk_call_up_date': f'Служба {service_base}:  {hospital_data.vk_urgent_call_date}' if hospital_data.vk_urgent_call_date else '',
        'personal_document' : patient.personal_document,
        'personal_number' : patient.personal_data,
        'anamnesis': hospital_data.anamnesis,
        'military_rank_p1' : get_accs_case(patient.military_rank.lower()) if get_accs_case(patient.military_rank.lower()) else patient.military_rank.lower(),
        'military_rank' : patient.military_rank.lower(),
        'ICD' : hospital_data.ICD,
        'date_of_illness' : patient.date_of_illness.strftime('%d {} %Y').format(month_name) if patient.date_of_illness else '',
        'address': patient.address,
        'branch': gent_branch if gent_branch else patient.branch.lower(),
        'therapist': hospital_data.therapist,
        'branch_v2': f'{branches_communications_add[patient.branch]} {get_gent_case(patient.branch).lower()}' if get_gent_case(patient.branch) else f'{branches_communications_add[patient.branch]} {patient.branch.lower()}',
        'name_v2': f'{branches_communications_rank[patient.branch]} м/с                                            {branches_communications_name[patient.branch]}' if branches_communications_rank[patient.branch] == 'капитан' or branches_communications_rank[patient.branch] == 'майор' else f'{branches_communications_rank[patient.branch]}                                            {branches_communications_name[patient.branch]}'
    }
    try:
        output_path = fill_template_v_times(template_path, context, patient.full_name, file_name='Представление на тяжесть')
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/vvk-report-to-category-g/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    template_path = os.path.join(docs_dir, "templates\\_Все для ввк на категорию годности г.docx")

    service_base = None
    if patient.service_basis == 'Военнослужащий по контракту':
        service_base = "по контракту"
    elif patient.service_basis == "Военнослужащий, призванный мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Военнослужащий, призванный по мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Курсант 1 курса (призывник)" or patient.service_basis == "Военнослужащий по призыву":
        service_base = "по призыву"
    else:
        service_base = "по контракту"

    
   
    patient_current_month = datetime.now().month
    month_name = month_names[patient_current_month]

    gent_branch = get_gent_case(patient.branch).lower()

    context = {
        'full_name': patient.full_name.upper(),
        'full_name_p2': patient.full_name,
        'today': datetime.now().strftime('%d {} %Y').format(month_name),
        'full_name_p1': get_accs_case(patient.full_name) if get_accs_case(patient.full_name).count == patient.full_name.count else patient.full_name,
        'branch_p': f'{branches_communications_add[patient.branch]} {gent_branch}' if gent_branch else patient.branch.lower(),
        'branch_rank': branches_communications_rank[patient.branch],
        'name_p1': branches_communications_name[patient.branch],
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_basis': service_base.lower() if service_base.lower() else service_base,
        'birthday': patient.birthday_date.strftime('%d.%m.%Y'),
        'finally_diagnos': hospital_data.expert_diagnosis,
        'vk_urgent_call_date': f'Служба по призыву: {hospital_data.vk_call_up_date}' if hospital_data.vk_call_up_date else 'Служба по призыву: не служил',
        'vk_call_up_date': f'Служба {service_base}:  {hospital_data.vk_urgent_call_date}' if hospital_data.vk_urgent_call_date else '',
        'personal_document' : patient.personal_document,
        'personal_number' : patient.personal_data,
        'anamnesis': hospital_data.anamnesis,
        'military_rank_p1' : get_accs_case(patient.military_rank.lower()) if get_accs_case(patient.military_rank.lower()) else patient.military_rank.lower(),
        'military_rank' : patient.military_rank.lower(),
        'ICD' : hospital_data.ICD,
        'date_of_illness' : patient.date_of_illness.strftime('%d {} %Y').format(month_name) if patient.date_of_illness else '',
        'address': patient.address,
        'branch': gent_branch if gent_branch else patient.branch.lower(),
        'therapist': hospital_data.therapist,
        'branch_v2': f'{branches_communications_add[patient.branch]} {get_gent_case(patient.branch).lower()}' if get_gent_case(patient.branch) else f'{branches_communications_add[patient.branch]} {patient.branch.lower()}',
        'name_v2': f'{branches_communications_rank[patient.branch]} м/с                                            {branches_communications_name[patient.branch]}' if branches_communications_rank[patient.branch] == 'капитан' or branches_communications_rank[patient.branch] == 'майор' else f'{branches_communications_rank[patient.branch]}                                            {branches_communications_name[patient.branch]}'
    }
    try:
        output_path = fill_template_v_times(template_path, context, patient.full_name, file_name='Представление на категорию Г')
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/vvk-report-to-category-v/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    template_path = os.path.join(docs_dir, "templates\\_Все для ввк на категорию годности в.docx")

    service_base = None
    if patient.service_basis == 'Военнослужащий по контракту':
        service_base = "по контракту"
    elif patient.service_basis == "Военнослужащий, призванный мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Военнослужащий, призванный по мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Курсант 1 курса (призывник)" or patient.service_basis == "Военнослужащий по призыву":
        service_base = "по призыву"
    else:
        service_base = "по контракту"


   
    patient_current_month = datetime.now().month
    month_name = month_names[patient_current_month]

    gent_branch = get_gent_case(patient.branch).lower()

    context = {
        'full_name': patient.full_name.upper(),
        'full_name_p2': patient.full_name,
        'today': datetime.now().strftime('%d {} %Y').format(month_name),
        'full_name_p1': get_accs_case(patient.full_name) if get_accs_case(patient.full_name).count == patient.full_name.count else patient.full_name,
        'branch_p': f'{branches_communications_add[patient.branch]} {gent_branch}' if gent_branch else patient.branch.lower(),
        'branch_rank': branches_communications_rank[patient.branch],
        'name_p1': branches_communications_name[patient.branch],
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_basis': service_base.lower() if service_base.lower() else service_base,
        'birthday': patient.birthday_date.strftime('%d.%m.%Y'),
        'finally_diagnos': hospital_data.expert_diagnosis,
        'vk_urgent_call_date': f'Служба по призыву: {hospital_data.vk_call_up_date}' if hospital_data.vk_call_up_date else 'Служба по призыву: не служил',
        'vk_call_up_date': f'Служба {service_base}:  {hospital_data.vk_urgent_call_date}' if hospital_data.vk_urgent_call_date else '',
        'personal_document' : patient.personal_document,
        'personal_number' : patient.personal_data,
        'anamnesis': hospital_data.anamnesis,
        'military_rank_p1' : get_accs_case(patient.military_rank.lower()) if get_accs_case(patient.military_rank.lower()) else patient.military_rank.lower(),
        'military_rank' : patient.military_rank.lower(),
        'ICD' : hospital_data.ICD,
        'date_of_illness' : patient.date_of_illness.strftime('%d {} %Y').format(month_name) if patient.date_of_illness else '',
        'address': patient.address,
        'branch': gent_branch if gent_branch else patient.branch.lower(),
        'therapist': hospital_data.therapist,
        'branch_v2': f'{branches_communications_add[patient.branch]} {get_gent_case(patient.branch).lower()}' if get_gent_case(patient.branch) else f'{branches_communications_add[patient.branch]} {patient.branch.lower()}',
        'name_v2': f'{branches_communications_rank[patient.branch]} м/с                                            {branches_communications_name[patient.branch]}' if branches_communications_rank[patient.branch] == 'капитан' or branches_communications_rank[patient.branch] == 'майор' else f'{branches_communications_rank[patient.branch]}                                            {branches_communications_name[patient.branch]}'
    }
    try:
        output_path = fill_template_v_times(template_path, context, patient.full_name, file_name='Представление на категорию В')
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/vvk-report-to-category-d/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    template_path = os.path.join(docs_dir, "templates\\_Все для ввк на категорию годности д.docx")

    service_base = None
    if patient.service_basis == 'Военнослужащий по контракту':
        service_base = "по контракту"
    elif patient.service_basis == "Военнослужащий, призванный мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Военнослужащий, призванный по мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Курсант 1 курса (призывник)" or patient.service_basis == "Военнослужащий по призыву":
        service_base = "по призыву"
    else:
        service_base = "по контракту"

   
   
    patient_current_month = datetime.now().month
    month_name = month_names[patient_current_month]

    gent_branch = get_gent_case(patient.branch).lower()

    context = {
        'full_name': patient.full_name.upper(),
        'full_name_p2': patient.full_name,
        'today': datetime.now().strftime('%d {} %Y').format(month_name),
        'full_name_p1': get_accs_case(patient.full_name) if get_accs_case(patient.full_name).count == patient.full_name.count else patient.full_name,
        'branch_p': f'{branches_communications_add[patient.branch]} {gent_branch}' if gent_branch else patient.branch.lower(),
        'branch_rank': branches_communications_rank[patient.branch],
        'name_p1': branches_communications_name[patient.branch],
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_basis': service_base.lower() if service_base.lower() else service_base,
        'birthday': patient.birthday_date.strftime('%d.%m.%Y'),
        'finally_diagnos': hospital_data.expert_diagnosis,
        'vk_urgent_call_date': f'Служба по призыву: {hospital_data.vk_call_up_date}' if hospital_data.vk_call_up_date else 'Служба по призыву: не служил',
        'vk_call_up_date': f'Служба {service_base}:  {hospital_data.vk_urgent_call_date}' if hospital_data.vk_urgent_call_date else '',
        'personal_document' : patient.personal_document,
        'personal_number' : patient.personal_data,
        'anamnesis': hospital_data.anamnesis,
        'military_rank_p1' : get_accs_case(patient.military_rank.lower()) if get_accs_case(patient.military_rank.lower()) else patient.military_rank.lower(),
        'military_rank' : patient.military_rank.lower(),
        'ICD' : hospital_data.ICD,
        'date_of_illness' : patient.date_of_illness.strftime('%d {} %Y').format(month_name) if patient.date_of_illness else '',
        'address': patient.address,
        'branch': gent_branch if gent_branch else patient.branch.lower(),
        'therapist': hospital_data.therapist,
        'branch_v2': f'{branches_communications_add[patient.branch]} {get_gent_case(patient.branch).lower()}' if get_gent_case(patient.branch) else f'{branches_communications_add[patient.branch]} {patient.branch.lower()}',
        'name_v2': f'{branches_communications_rank[patient.branch]} м/с                                            {branches_communications_name[patient.branch]}' if branches_communications_rank[patient.branch] == 'капитан' or branches_communications_rank[patient.branch] == 'майор' else f'{branches_communications_rank[patient.branch]}                                            {branches_communications_name[patient.branch]}'
    }
    try:
        output_path = fill_template_v_times(template_path, context, patient.full_name, file_name='Представление на категорию Д')
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/vvk-conclusion/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    service_base = None
    if patient.service_basis == 'Военнослужащий по контракту':
        service_base = "По контракту"
    elif patient.service_basis == "Военнослужащий, призванный мобилизации":
        service_base = "По мобилизации"
    elif patient.service_basis == "Курсант 1 курса (призывник)" or patient.service_basis == "Военнослужащий по призыву":
        service_base = "По призыву"
    else:
        service_base = "По контракту"
   
    patient_current_month = datetime.now().month
    month_name = month_names[patient_current_month]

    template_path = os.path.join(docs_dir, "templates\\Заключение ВВК шаблон.docx")
    context = {
        'full_name': patient.full_name.upper(),
        'date_today': datetime.now().strftime('%d {} %Y г.').format(month_name),
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_basis': service_base,
        'birhday': patient.birthday_date.strftime('%d.%m.%Y'),
        'final_diagnosis': hospital_data.expert_diagnosis,
        'vk_urgent_call_date': f'Служба по призыву: {hospital_data.vk_urgent_call_date}' if hospital_data.vk_urgent_call_date else 'Служба по призыву: не служил',
        'vk_call_up_date': hospital_data.vk_call_up_date,
        'personal_document' : patient.personal_document,
        'personal_number' : patient.personal_document,
        'anamnesis': hospital_data.anamnesis,
        'military_rank' : patient.military_rank.lower(),
        'therapist': hospital_data.therapist
    }
    try:
        output_path = fill_template(template_path, context, patient.full_name, file_name=f'Заключение ВВК {patient.id}')
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/about-finding/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):

    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")
    




    additionals = additional[patient.branch]

    ablt_branch = get_loct_case(patient.branch)



    template_path = os.path.join(docs_dir, "templates\\Справка о нахождении на госпитализации шаблон.docx")
    context = {
        'full_name': patient.full_name if patient.full_name else patient.full_name,
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_basis': patient.service_basis,
        'birthday': patient.birthday_date.strftime('%d.%m.%Y') if patient.birthday_date.strftime('%d.%m.%Y') else None,
        'final_diagnosis': hospital_data.final_diagnosis,
        'vk_urgent_call_date': hospital_data.vk_urgent_call_date,
        'vk_call_up_date': hospital_data.vk_call_up_date,
        'personal_document' : patient.personal_document,
        'anamnesis': hospital_data.anamnesis,
        'military_rank' :patient.military_rank.lower() if patient.military_rank.lower() else patient.military_rank.lower(),
        'hospitalization_date': patient.current_time.strftime('%d.%m.%Y'),
        'branch' : ablt_branch if ablt_branch + additionals else patient.branch.lower() + additionals,
        'history_number': patient.history_number,
        'personal_number': patient.personal_data
    }
    try:
        output_path = fill_template(template_path, context, patient.full_name, file_name='Справка о нахождении на госпитализации')
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/list-of-hospitals-download/{patient_id}")
async def generate_document(request: Request):
    try:        
        output_path = os.path.join(docs_dir, f'Список больниц в эпикризы.docx')
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/injury-severity-report/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    template_path = os.path.join(docs_dir, "templates\\Справка о ранении образец.docx")

       
    patient_current_month = patient.current_time.month
    month_name = month_names[patient_current_month]

    month_name = month_name.replace('"', '')

    service_base = None
    if patient.service_basis == 'Военнослужащий по контракту':
        service_base = "по контракту"
    elif patient.service_basis == "Военнослужащий, призванный мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Военнослужащий, призванный по мобилизации":
        service_base = "по мобилизации"
    elif patient.service_basis == "Курсант 1 курса (призывник)" or patient.service_basis == "Военнослужащий по призыву":
        service_base = "по призыву"
    else:
        service_base = "по контракту"

    context = {
        'full_name': patient.full_name,
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_base': service_base,
        'birthday': patient.birthday_date.strftime('%d.%m.%Y') if patient.birthday_date.strftime('%d.%m.%Y') else None,
        'final_diagnosis': hospital_data.diagnosis_according_to_form_one_hundred,
        'date_of_injury': patient.date_of_illness.strftime('%d.%m.%Y') if patient.date_of_illness else None,
        'date_of_hospitalization': patient.current_time.strftime('%d {} %Y г.').format(month_name),
        'military_rank' : patient.military_rank
    }

    try:
        output_path = fill_template(template_path, context, patient.full_name)
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/discharge-epieriosis/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    movements = db.query(models.PatientMovement)\
                .filter_by(patient_id=patient.id)\
                .order_by(models.PatientMovement.event_date.desc())\
                .first()

     # Преобразуем даты в объекты datetime
    date_of_hospitalization = patient.current_time
    date_of_discharge = movements.event_date or datetime.now()

    # Вычисляем разницу и берем абсолютное значение
    delta = abs((date_of_discharge - date_of_hospitalization).days) + 1

    

    template_path = os.path.join(docs_dir, "templates\_Выписной эпикриз шаблон.docx")
    
    context = {
        'flag_on' : 'оформлена' if hospital_data.certificate_of_injury else 'не оформлена',
        'history_number': patient.history_number,
        'full_name_case': format_name(patient.full_name),
        'current_time': date_of_discharge.strftime('%d.%m.%Y') or date_of_discharge.strftime('%d.%m.%Y') or datetime.now('%d.%m.%Y'),
        'patient_id': patient.history_number, 
        'full_name': patient.full_name,
        'date_of_discharge': date_of_discharge.strftime('%d.%m.%Y') or date_of_discharge.strftime('%d.%m.%Y') or datetime.now('%d.%m.%Y'),
        'now_year': datetime.now().year,
        'delta': delta,
        'military_unit': patient.military_unit,
        'service_base': patient.service_basis,
        'birhday': patient.birthday_date.strftime('%d.%m.%Y'),
        'final_diagnosis': hospital_data.final_diagnosis,
        'date_of_injury': patient.date_of_illness.strftime('%d.%m.%Y') if patient.date_of_illness else None,
        'date_of_hospitalization': date_of_hospitalization.strftime('%d.%m.%Y'),
        'military_rank' : patient.military_rank,
        'anamines': hospital_data.anamnesis,
        'today' : datetime.now().strftime('%d.%m.%Y'),
        'address': patient.address,
        'branch': get_loct_case(patient.branch.lower()) if get_loct_case(patient.branch.lower()) else patient.branch.lower(),
        'name' : branches_communications_name[patient.branch],
        'therapist': hospital_data.therapist
    }
    try:
        output_path = fill_template(template_path, context, patient.full_name, file_name='Выписной эпикриз' )
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/reference-report/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    movements = db.query(models.PatientMovement)\
                .filter_by(patient_id=patient.id)\
                .order_by(models.PatientMovement.event_date.desc())\
                .first()

    certificate_of_injury = db.query(models.CertificateOfInjury)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.CertificateOfInjury.load_date.desc())\
        .first()

    certificate_of_severity = db.query(models.CertificateOfSeverity)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.CertificateOfSeverity.severity_date.desc())\
        .first()
    
    vvk = db.query(models.VvkHospital)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.VvkHospital.id.desc())\
        .first()
    
     # Преобразуем даты в объекты datetime
    date_of_hospitalization = patient.current_time
    date_of_discharge = movements.event_date or datetime.now()

    # Вычисляем разницу и берем абсолютное значение
    delta = abs((date_of_discharge - date_of_hospitalization).days) + 1

    try:
        vvk_date = vvk.vvk_date.strftime('%d.%m.%Y')
    except:
        vvk_date = None
   
    try:
       certificate_of_severity_t = certificate_of_severity.approval_number
    except:
        certificate_of_severity_t = None

    try:
        certificate_of_injury_t = certificate_of_injury.injury_number
    except: 
        certificate_of_injury_t = None

    

    template_path = os.path.join(docs_dir, "templates\Справка доклад на пациента.docx")
    
    context = {
        'Фамилия и инициалы': patient.full_name,
        'Номер военного билета': patient.personal_document,
        'Отделение': patient.branch,
        'Номер телефона': patient.phone_number if patient.phone_number else 'не был зафиксирован',
        'Адрес': patient.address,
        'Военный комиссариат и дата призыва на службу': hospital_data.vk_urgent_call_date,
        'Округ': hospital_data.district,
        'Кем направлен': patient.directions.lower(),
        'Диагноз при поступлении': patient.first_diagnosis,
        'Лечащий врач': hospital_data.therapist,
        'Первичный диагноз': hospital_data.diagnosis_upon_admission,
        'Анамнез': hospital_data.anamnesis if hospital_data.anamnesis else '(не зафиксирован)',
        'Экспертный диагноз': hospital_data.expert_diagnosis if hospital_data.expert_diagnosis else '(не зафиксирован)',
        'Окончательный диагноз': hospital_data.final_diagnosis if hospital_data.final_diagnosis else '(не зафиксирован)',
        'Личный номер': patient.personal_data if patient.personal_data else '(не зафиксирован)',
        'Дата госпитализации': patient.current_time.strftime('%d.%m.%Y'),
        'Дата последнего события': movements.event_date.strftime('%d.%m.%Y'),
        'Событие': movements.event_type.lower(),
        'Примечание': movements.note if movements.note else '(отсутсвует)' ,
        'Основа службы': patient.service_basis,
        'Дата выдачи ВВК': f', дата выдачи заключения ввк {vvk_date}' if vvk_date else '',  
        'Номер заключения ВВК': f'Номер заключения ВВК {vvk.vvk_number}' if vvk_date else 'Заключения ВВК не зафиксированы.',
        'Заключение ВВК': f', заключение ВВК: {vvk.conclusion}' if vvk_date else '',

        'Номер справки подтверждения': f'Была подтверждена отделом ВВЭ под № {certificate_of_severity.approval_number}' if certificate_of_severity_t  else '' ,
        'Дата выдачи справки о подтверждении': f'{certificate_of_severity.approval_date.strftime("%d.%m.%Y")}' if certificate_of_severity_t else '' ,
        'Номер справки тяжести увечья': certificate_of_severity.approval_date if certificate_of_severity_t else '',
        'Дата справки о тяжести увечья': certificate_of_severity.approval_date if certificate_of_severity_t  else 'не выдавалась, в связи с тем, что справка не оформлялась.',
        'Номер справки о ранении': certificate_of_injury.injury_number if certificate_of_injury_t else '',
        'Дата выдачи справки о ранении': certificate_of_injury.certificate_injury_date if certificate_of_injury_t else 'не выданна.',
        'Должность': branches_communications_add[patient.branch] + ' ' + branch_mapping[patient.branch],
        'Звание должностного лица': branches_communications_rank[patient.branch],
        'ФИО': branches_communications_name[patient.branch],        
        'Воинская часть': patient.military_unit,
        'Дата рождения': patient.birthday_date.strftime('%d.%m.%Y'),
        'Звание' : patient.military_rank.lower(),
    }
    try:
        output_path = fill_template(template_path, context, patient.full_name, file_name='Справка доклад на' )
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/extract-report/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    movements = db.query(models.PatientMovement)\
                .filter_by(patient_id=patient.id)\
                .order_by(models.PatientMovement.event_date.desc())\
                .first()

    certificate_of_injury = db.query(models.CertificateOfInjury)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.CertificateOfInjury.load_date.desc())\
        .first()

    certificate_of_severity = db.query(models.CertificateOfSeverity)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.CertificateOfSeverity.severity_date.desc())\
        .first()
    
    vvk = db.query(models.VvkHospital)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.VvkHospital.id.desc())\
        .first()
    
     # Преобразуем даты в объекты datetime
    date_of_hospitalization = patient.current_time
    date_of_discharge = movements.event_date or datetime.now()

    # Вычисляем разницу и берем абсолютное значение
    delta = abs((date_of_discharge - date_of_hospitalization).days) + 1

    try:
        vvk_date = vvk.vvk_date.strftime('%d.%m.%Y')
    except:
        vvk_date = None
   
    try:
       certificate_of_severity_t = certificate_of_severity.approval_number
    except:
        certificate_of_severity_t = None

    try:
        certificate_of_injury_t = certificate_of_injury.injury_number
    except: 
        certificate_of_injury_t = None

    

    template_path = os.path.join(docs_dir, "templates\Уведомление о выписке шаблон ВГ.docx")
    
    context = {
        'Фамилия и инициалы': patient.full_name,
        'Направительный диагноз': patient.first_diagnosis,
        'Дата госпитализации': patient.current_time.strftime('%d.%m.%Y'),
        'Дата выписки': movements.event_date.strftime('%d.%m.%Y'),
        'Основа службы': patient.service_basis,
        'Куда': movements.note,
        'Должность': "Временно исполняющий обязанности начальника",
        'Имя подписывающего': 'Р. Сулейманов',
        'Звание должностного лица': 'подполковник медицинской службы',
        'ФИО': branches_communications_name[patient.branch],        
        'Воинская часть': patient.military_unit,
        'Дата рождения': patient.birthday_date.strftime('%d.%m.%Y'),
        'Звание' : patient.military_rank.lower(),
    }
    try:
        output_path = fill_template(template_path, context, patient.full_name, file_name='Уведомление о выписке ' )
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/at-hostpitalization-report/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    movements = db.query(models.PatientMovement)\
                .filter_by(patient_id=patient.id)\
                .order_by(models.PatientMovement.event_date.desc())\
                .first()

    certificate_of_injury = db.query(models.CertificateOfInjury)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.CertificateOfInjury.load_date.desc())\
        .first()

    certificate_of_severity = db.query(models.CertificateOfSeverity)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.CertificateOfSeverity.severity_date.desc())\
        .first()
    
    vvk = db.query(models.VvkHospital)\
        .filter_by(patient_id=patient_id)\
        .order_by(models.VvkHospital.id.desc())\
        .first()
    
     # Преобразуем даты в объекты datetime
    date_of_hospitalization = patient.current_time
    date_of_discharge = movements.event_date or datetime.now()

    # Вычисляем разницу и берем абсолютное значение
    delta = abs((date_of_discharge - date_of_hospitalization).days) + 1

    try:
        vvk_date = vvk.vvk_date.strftime('%d.%m.%Y')
    except:
        vvk_date = None
   
    try:
       certificate_of_severity_t = certificate_of_severity.approval_number
    except:
        certificate_of_severity_t = None

    try:
        certificate_of_injury_t = certificate_of_injury.injury_number
    except: 
        certificate_of_injury_t = None

    

    template_path = os.path.join(docs_dir, "templates\Уведомление о нахождении на госпитализации шаблон ВГ.docx")
    
    context = {
        'ФИО': patient.full_name,
        'Направительный диагноз': patient.first_diagnosis,
        'Дата госпитализации': patient.current_time.strftime('%d.%m.%Y'),
        'Должностное лицо': "Временно исполняющий обязанности начальника",
        'ФИО должностного лица': 'Р. Сулейманов',
        'Звание должностного лица': 'подполковник медицинской службы',
        'Войсковая часть': patient.military_unit,
        'Дата рождения': patient.birthday_date.strftime('%d.%m.%Y'),
        'Звание' : patient.military_rank.lower(),
    }
    try:
        output_path = fill_template(template_path, context, patient.full_name, file_name='Уведомление о нахождении на госпитализации ' )
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/wilting-direction/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    template_path = os.path.join(docs_dir, "templates\\Заявка направление шаблон.docx")
    context = {
        'full_name': patient.full_name,
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_base': patient.service_basis,
        'birhday': patient.birthday_date.strftime('%d.%m.%Y'),
        'final_diagnosis': hospital_data.final_diagnosis,
        'date_of_injury': patient.date_of_illness.strftime('%d.%m.%Y') if patient.date_of_illness else None,
        'date_of_hospitalization': patient.current_time.strftime('%d.%m.%Y'),
        'military_rank' : patient.military_rank,
        'anamines': hospital_data.anamnesis,
        'today' : datetime.now().strftime('%d.%m.%Y'),
        'address': patient.address,
        'branch': patient.branch,
        'therapist': hospital_data.therapist
    }
    try:
        output_path = fill_template(template_path, context, patient.full_name, file_name='Заявка, направление' )
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/raport-new/{patient_id}")
async def generate_document(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    patient = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    hospital_data = db.query(models.HospitalData).filter(models.HospitalData.patient_id == patient_id).first()
    if not hospital_data:
        raise HTTPException(status_code=404, detail="Hospital data not found")

    birhday = patient.birthday_date.strftime('%d.%m.%Y')

    
    patient_current_month = datetime.now().month
    month_name = month_names[patient_current_month]

    month_name = month_name.replace('"', '')


    branch_name = branches_communications[patient.branch]

    template_path = os.path.join(docs_dir, "templates\\Новый Образец Рапорта Сулейманов.docx")
    context = {
        'day': datetime.now().strftime('%d'),
        'branch_full_desc': f'{branches_communications_add[patient.branch]} {get_gent_case(patient.branch).lower()}' if get_dative_case(branches_communications_add[patient.branch]) or get_gent_case(patient.branch) else f'Начальник {patient.branch.lower()}',
        'branch_full_desc_v':  f'{get_dative_case(branches_communications_add[patient.branch])} {get_gent_case(patient.branch).lower()}' if branches_communications_add[patient.branch] == "Начальник" else f'{branches_communications_add[patient.branch]} {get_gent_case(patient.branch).lower()}',
        'month': month_names[int(datetime.now().strftime('%m'))] if month_names[int(datetime.now().strftime('%m'))] else datetime.now().strftime('%m'),
        'year': datetime.now().strftime('%Y'),
        'currnet_time': patient.current_time.strftime('%d.%m.%Y'),
        'full_name': get_dative_case(patient.full_name) if get_dative_case(patient.full_name) else patient.full_name ,
        'now_year': datetime.now().year,
        'military_unit': patient.military_unit,
        'service_base': patient.service_basis,
        'personal_number': patient.personal_data,
        'history_number': patient.history_number,
        'diagnos_form_100': f'{hospital_data.diagnosis_according_to_form_one_hundred}. Дата получения ранения {patient.date_of_illness.strftime("%d.%m.%Y")} г.',
        'birthday': birhday if birhday else None,
        'final_diagnosis': hospital_data.final_diagnosis,
        'date_of_injury': patient.date_of_illness.strftime('%d.%m.%Y') if patient.date_of_illness else None,
        'date_of_hospitalization': patient.current_time.strftime('%d.%m.%Y'),
        'military_rank' : patient.military_rank.lower(),
        'anamines': hospital_data.anamnesis,
        'today' : datetime.now().strftime('%d.%m.%Y'),
        'address': patient.address,
        'branch': get_loct_case(patient.branch).lower(),
        'branch_name': branch_name,
        'therapist': hospital_data.therapist
    }
    try:
        output_path = fill_template_v_times(template_path, context, patient.full_name, file_name='Рапорт на справку о ранении' )
        download_link = request.url_for('download_report', filename=os.path.basename(output_path))
        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, output_path, 300)
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": str(download_link)})
    except Exception as e:
        logger.error(f"Error while generating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/download_report/{filename}", response_class=FileResponse)
def download_report(filename: str, background_tasks: BackgroundTasks):
    file_path = os.path.join(docs_dir, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    # Добавляем задачу удаления файла через 5 минут после его загрузки
    delete_file_with_delay(file_path, 300, background_tasks)

    return FileResponse(path=file_path, filename=filename)

@router.get("/create_patient_diary/{patient_id}", tags=["Patient Diaries"])
async def create_patient_diary_endpoint(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    try:
        # Получение данных о пациенте
        patient_data = db.query(models.PassportData).filter(models.PassportData.id == patient_id).first()
        if not patient_data:
            raise HTTPException(status_code=404, detail="Patient not found")
        
        # Создание дневника пациента
        doc_path = create_patient_diary(patient_data, db)

        download_link = request.url_for('download_report', filename=os.path.basename(doc_path))
        download_link = str(download_link)  # Making sure the link is a string

        # Добавление задачи на удаление файла через 5 минут
        background_tasks.add_task(delete_file_with_delay, doc_path, 300)

        return JSONResponse(content={"message": "Report generated successfully.", "download_link": download_link})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Первичка
@router.get("/generate_medical_report/{patient_id}")
async def generate_medical_report(request: Request, patient_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    try:
        doc_path = create_medical_report(patient_id, db)
        
        download_link = request.url_for('download_report', filename=os.path.basename(doc_path))
        download_link = str(download_link)  # Making sure the link is a string

        delete_file_with_delay(doc_path, 300, background_tasks)

        return JSONResponse(content={"message": "Report generated successfully.", "download_link": download_link})
    except HTTPException as e:
        return JSONResponse(status_code=e.status_code, content={"detail": e.detail})
    except Exception as e:
        logger.error(f"Error generating medical report: {e}")
        return JSONResponse(status_code=500, content={"detail": f"Internal server error: {str(e)}"})

#endregion

#region hello_page 

@router.get("/")
async def department_statistics(request: Request, db: Session = Depends(get_db)):
    now = datetime.now()
    if now.hour < 8:
        start_of_today = datetime(now.year, now.month, now.day, 0) 
    else:
        start_of_today = datetime(now.year, now.month, now.day, 0)
    end_of_today = start_of_today + timedelta(days=1)

    # Dictionary mapping departments to the number of beds they have
    bed_counts = {
        'Хирургическое отделение': 50,
        'Терапевтическое отделение': 40,
        'Неврологическое отделение': 25,
        'Кожно-венерологическое отделение': 25,#chhh
        'Инфекционное отделение': 30
    }

    # Subquery for the latest movements
    latest_movements = (
    db.query(
        models.PatientMovement.patient_id,
        func.max(models.PatientMovement.event_date).label('latest_date')
    )
    .group_by(models.PatientMovement.patient_id)
    .subquery())


    # Main query
    department_query = (
    db.query(
        models.PassportData.branch.label('department'),  # Используем branch как department
        func.count().filter(models.PatientMovement.event_type.in_([
                "Прием на госпитализацию",
                "Перевод в другое отделение",
                "Планируется перевод"
        ])).label('active_patients'),
        func.count().filter(and_(
            models.PatientMovement.event_type == 'Прием на госпитализацию',
            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today
        )).label('admitted'),
        func.count().filter(and_(
            or_(
                models.PatientMovement.event_type == 'Выписан',
                models.PatientMovement.event_type == 'Выписан за нарушение режима',
                models.PatientMovement.event_type == 'Направлен в санаторий',
            ),

            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today
        )).label('discharged'),
        func.count().filter(and_(
            models.PatientMovement.event_type == 'Перевод в другое отделение',
            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today,
        )).label('transferred_out'),
        func.count().filter(and_(
            models.PatientMovement.event_type == 'Перевод в другое отделение',
            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today
        )).label('transferred_in'),
        func.count().filter(and_(
            or_(
                models.PatientMovement.event_type == 'Перевод в другое МО',
                models.PatientMovement.event_type == 'Перевод в другое ВМО',
                models.PatientMovement.event_type == 'Перевод в РКБ',
            ),
            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today
        )).label('transferred_to_other_MO')
    )
    .join(models.PassportData, models.PatientMovement.patient_id == models.PassportData.id)  # Объединение с PassportData
    .join(latest_movements, and_(
        models.PatientMovement.patient_id == latest_movements.c.patient_id,
        models.PatientMovement.event_date == latest_movements.c.latest_date
    )) 
    .group_by(models.PassportData.branch))

    departments = department_query.all()

    # Preparing data for the template with bed counts
    departments_data = [
        {
            "department": dept.department,
            "beds_deployed": bed_counts.get(dept.department, 0),
            "active_patients": dept.active_patients,
            "admitted": dept.admitted,
            "discharged": dept.discharged,
            "transferred_out": dept.transferred_out,
            "transferred_in": dept.transferred_in,
            "transferred_to_other_MO": dept.transferred_to_other_MO,
            "current_patients": dept.active_patients,
            "factual_number": dept.active_patients,
            "load_percentage": (dept.active_patients / bed_counts.get(dept.department, 1)) * 100,
            "free_beds": bed_counts.get(dept.department, 0) - dept.active_patients
        }
        for dept in departments
    ]


    totals = {
        'beds_deployed': sum(dept['beds_deployed'] for dept in departments_data),
        'active_patients': sum(dept['active_patients'] for dept in departments_data),
        'admitted': sum(dept['admitted'] for dept in departments_data),
        'discharged': sum(dept['discharged'] for dept in departments_data),
        'transferred_out': sum(dept['transferred_out'] for dept in departments_data),
        'transferred_in': sum(dept['transferred_in'] for dept in departments_data),
        'transferred_to_other_MO': sum(dept['transferred_to_other_MO'] for dept in departments_data),
        'current_patients': sum(dept['current_patients'] for dept in departments_data),
        'factual_number': sum(dept['factual_number'] for dept in departments_data),
        'load_percentage': sum(dept['load_percentage'] for dept in departments_data) / len(departments_data) if departments_data else 0,  # Average load percentage
        'free_beds': sum(dept['free_beds'] for dept in departments_data)
    }

    patients_with_pneumonia = db.query(models.PassportData).join(models.HospitalData).filter(
        models.HospitalData.diagnosis_upon_admission.ilike("%пневмония%")
    ).order_by(models.PassportData.id.desc()).all()

    military_ranks_query = (
        db.query(
            models.PassportData.military_rank,
            func.count(models.PassportData.id).label('count')
        )
        .group_by(models.PassportData.military_rank)
    )

    military_ranks_data = military_ranks_query.all()

    today = datetime.now().date()
    start_of_month = today.replace(day=1)
    
    next_month = today.replace(day=28) + timedelta(days=4)
    end_of_month = next_month - timedelta(days=next_month.day)

    # Измененный запрос daily_statistics_query для месячной статистики
    daily_statistics_query = (
        db.query(
            func.date(models.PatientMovement.event_date).label('event_day'),
            func.count().filter(
                and_(
                    models.PatientMovement.event_type.in_([
                        'Прием на госпитализацию'
                    ]),
                    models.PatientMovement.event_date >= start_of_month,
                    models.PatientMovement.event_date <= end_of_month
                )
            ).label('count_v1'),
            func.count().filter(
                and_(
                    models.PatientMovement.event_type.in_([
                        'Выписан',
                        'Перевод в другое отделение',
                        'Перевод в другое ВМО',
                        'Перевод в другое МО',
                        'Перевод в РКБ',
                        'Направлен в санаторий',
                        'СОЧ',
                        'Летальный исход'
                    ]),
                    models.PatientMovement.event_date >= start_of_month,
                    models.PatientMovement.event_date <= end_of_month
                )
            ).label('count_v2'),
            # Другие агрегатные функции для статистики
        )
        .group_by(func.date(models.PatientMovement.event_date))
        .order_by(func.date(models.PatientMovement.event_date))
    )


    # Получение статистики за месяц
    daily_statistics = daily_statistics_query.all()

    chart_data = {
        "labels": [str(row.event_day) for row in daily_statistics],
        "counts": [row.count_v1 for row in daily_statistics],
        "counts_out": [row.count_v2 for row in daily_statistics],
        # Другие данные для графика, если необходимо
    }

    branch_counts_query = (
        db.query(
            models.PassportData.branch,
            func.count().label('filled_records'),
            func.count().label('total_records')
        )
        .group_by(models.PassportData.branch)
    )
    branch_counts_data = branch_counts_query.all()

    # Подготовка данных для графика
    branch_labels = [branch.branch for branch in branch_counts_data]
    filled_records = [branch.filled_records for branch in branch_counts_data]
    total_records = [branch.total_records for branch in branch_counts_data]
    percentage_filled = [(filled / total) * 100 if total != 0 else 0 for filled, total in zip(filled_records, total_records)]

    # Добавление данных о заполненности по отделениям в словарь для передачи в шаблон HTML
    branch_chart_data = {
        "labels": branch_labels,
        "filled_records": filled_records,
        "percentage_filled": percentage_filled
    }

    current_time = datetime.now()
    first_day_of_month = current_time.replace(day=1)

    

    passport_data_list = db.query(models.PassportData).filter(
        models.PassportData.current_time >= first_day_of_month,
        models.PassportData.current_time <= current_time
    ).all()

    branch_completeness = check_completeness_by_branch(passport_data_list)


    # Prepare data for the chart
    branch_completeness_chart_data = {
        "labels": list(branch_completeness.keys()),
        "completeness": list(branch_completeness.values())
    }
    
 
    referred_by_query = (
        db.query(
            models.PassportData.directions,
            func.count(models.PassportData.id).label('count')
        )
        .group_by(models.PassportData.directions)
    )

    referred_by_data = referred_by_query.all()

    referred_by_chart_data = {
        "labels": [row.directions for row in referred_by_data],
        "counts": [row.count for row in referred_by_data]
    }

    service_basis_query = (
        db.query(
            models.PassportData.service_basis,
            func.count(models.PassportData.id).label('count')
        )
        .group_by(models.PassportData.service_basis)
    )

    service_basis_data = service_basis_query.all()

    service_basis_labels = [record.service_basis for record in service_basis_data]
    service_basis_counts = [record.count for record in service_basis_data]


    # Получаем текущую дату
    current_year = datetime.now().year

    # Дата начала текущего года
    start_of_year = datetime(current_year, 1, 1)

    # Получаем данные по ICD за текущий год
    icd_query = (
        db.query(
            models.HospitalData.ICD,
            func.count(models.HospitalData.id).label('count')
        )
        .join(models.PatientMovement, models.PatientMovement.patient_id == models.HospitalData.patient_id)
        .filter(
            and_(
                models.PatientMovement.event_type == 'Прием на госпитализацию',
                models.PatientMovement.event_date >= start_of_year
            )
        )
        .distinct(models.HospitalData.ICD) 
        .group_by(models.HospitalData.ICD)
    )

    icd_data = icd_query.all()

    # Создаем список меток и количества для графика
    icd_labels = [record.ICD if record.ICD else 'Не проставлено' for record in icd_data]
    icd_counts = [record.count for record in icd_data] 

    # Получаем данные по воинским частям за текущий год
    military_units_query = (
        db.query(
            models.PassportData.military_unit,
            func.count().label('count')
        )
        .join(models.PatientMovement, models.PatientMovement.patient_id == models.PassportData.id)
        .filter(
            and_(
                models.PatientMovement.event_type.in_(['Прием на госпитализацию']),
                models.PatientMovement.event_date >= start_of_year
            )
        )
        .group_by(models.PassportData.military_unit)
    )


    military_units_data = military_units_query.all()

    # Подготовка данных для графика
    military_units_labels = [record.military_unit if record.military_unit else 'Не проставлено' for record in military_units_data]
    military_units_counts = [record.count for record in military_units_data]

    therapist_query = (
        db.query(
            models.HospitalData.therapist,
            func.count(models.HospitalData.id).label('count')
        )
        .group_by(models.HospitalData.therapist)
    )


    therapist_data = therapist_query.all()

    # Подготовка данных для графика
    therapist_labels = [record.therapist if record.therapist else 'Не проставлено' for record in therapist_data]
    therapist_counts = [record.count for record in therapist_data]

    # Получение данных о перемещениях пациентов по месяцам
    movements = db.query(
        models.PatientMovement.department,
        extract('month', models.PatientMovement.event_date).label('month'),
        func.count(models.PatientMovement.id).label('count')
    ).filter(
        models.PatientMovement.event_type == 'Прием на госпитализацию'
    ).group_by(
        models.PatientMovement.department,
        extract('month', models.PatientMovement.event_date)
    ).all()

    # Преобразование данных в нужный формат
    monthly_department_data = {
        'labels': ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь'],
        'data': [],
        'total_counts': [0] * 12
    }

    # Организация данных по отделениям и месяцам
    department_data = {}
    for movement in movements:
        department = movement.department
        month = int(movement.month) - 1  # Преобразование месяца в индекс (0-11)
        count = movement.count
        
        if department not in department_data:
            department_data[department] = [0] * 12
        department_data[department][month] += count
        monthly_department_data['total_counts'][month] += count

    # Заполнение данных по отделениям
    for department, counts in department_data.items():
        monthly_department_data['data'].append({
            'department': department,
            'counts': counts
        })

    #Койки и т.д.

    # Получение данных о движениях пациентов
    patient_movements = db.query(
        models.PatientMovement.patient_id,
        models.PatientMovement.event_type,
        models.PatientMovement.event_date
    ).filter(
        models.PatientMovement.event_type.in_([
            'Прием на госпитализацию', 'Выписан', 'Выписан за нарушение режима',
            'Перевод в другое ВМО', 'Направлен в санаторий', 'Перевод в другое МО',
            'Перевод в РКБ', 'СОЧ', 'Летальный исход'
        ])
    ).order_by(models.PatientMovement.patient_id, models.PatientMovement.event_date).all()

    # Обработка данных для вычисления метрик
    hospitalization_days = {}  # Словарь для хранения данных о днях госпитализации
    total_days_per_month = {}  # Словарь для хранения суммарного количества дней за каждый месяц

    for record in patient_movements:
        if record.event_type == 'Прием на госпитализацию':
            admission_date = record.event_date
        elif record.event_type in ['Выписан', 'Выписан за нарушение режима', 'Перевод в другое ВМО', 'Направлен в санаторий', 'Перевод в другое МО', 'Перевод в РКБ', 'СОЧ', 'Летальный исход']:
            discharge_date = record.event_date
            if admission_date and discharge_date:
                month = admission_date.month
                year = admission_date.year
                days_stayed = (discharge_date - admission_date).days

                if (year, month) not in hospitalization_days:
                    hospitalization_days[(year, month)] = []

                hospitalization_days[(year, month)].append(days_stayed)

                if (year, month) not in total_days_per_month:
                    total_days_per_month[(year, month)] = 0

                total_days_per_month[(year, month)] += days_stayed
                admission_date = None  # Сброс даты поступления для следующего пациента

   # Вычисление метрик
    bed_turnover = {}  # Оборот койки
    bed_utilization = {}  # Коечная мощность
    average_length_of_stay = {}  # Средняя длительность пребывания

    total_beds = 100  # Предполагаемое количество коек в больнице

    for key, days in hospitalization_days.items():
        year, month = key
        total_days = sum(days)
        num_patients = len(days)

        bed_turnover[key] = total_days_per_month[key]
        bed_utilization[key] = (total_days_per_month[key] / (30 * total_beds)) * 100  # Процент использования коек
        average_length_of_stay[key] = total_days / num_patients

    # Преобразование данных в нужный формат для передачи в шаблон
    metrics_data = {
        'labels': ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь'],
        'data': {
            'bed_turnover': bed_turnover,
            'bed_utilization': bed_utilization,
            'average_length_of_stay': average_length_of_stay
        }
    }


    return templates.TemplateResponse("department_statistics.html", {
        "request": request,
        "departments": departments_data,
        "patients_with_pneumonia" : patients_with_pneumonia,
        "military_ranks_data" :  military_ranks_data,
        "totals": totals,
        "title": "Статистика по отделениям",
        "chart_data": chart_data,
        "branch_chart_data": branch_chart_data,
        "branch_completeness_chart_data": branch_completeness_chart_data,
        "referred_by_chart_data": referred_by_chart_data,
        "service_basis_labels" : service_basis_labels,
        "service_basis_counts" : service_basis_counts,
        "icd_labels": icd_labels,
        "icd_counts": icd_counts,
        "military_units_labels": military_units_labels,
        "military_units_counts": military_units_counts,
        "therapist_labels": therapist_labels,
        "therapist_counts": therapist_counts,
        "monthly_department_data": monthly_department_data,
        "metrics_data": metrics_data

    })

@router.get("/api/hospitalization_data")
async def get_hospitalization_data(month: str, db: Session = Depends(get_db)):
    """
    Получает данные о госпитализациях и выписках за указанный месяц.
    
    Args:
        month (str): Месяц в формате "ГГГГ-ММ".
        db (Session): Сессия базы данных.
    
    Returns:
        dict: Данные для графика.
    """
    start_date = datetime.strptime(month, "%Y-%m")
    end_date = start_date + timedelta(days=30)

    daily_statistics_query = (
        db.query(
            func.date(models.PatientMovement.event_date).label('event_day'),
            func.sum(case((models.PatientMovement.event_type == "Прием на госпитализацию", 1), else_=0)).label('count_v1'),
            func.sum(case((models.PatientMovement.event_type == "Выписан", 1), else_=0)).label('count_v2')
        )
        .filter(models.PatientMovement.event_date >= start_date, models.PatientMovement.event_date < end_date)
        .group_by(func.date(models.PatientMovement.event_date))
        .order_by(func.date(models.PatientMovement.event_date))
    )

    daily_statistics = daily_statistics_query.all()

    chart_data = {
        "labels": [str(row.event_day) for row in daily_statistics],
        "counts": [row.count_v1 for row in daily_statistics],
        "counts_out": [row.count_v2 for row in daily_statistics],
    }

    return chart_data


@router.get("/download_department_statistics/")
async def download_department_statistics(request: Request, db: Session = Depends(get_db)):
    # Get the current date and time
    now = datetime.now()

    # Define the start and end of the previous day
    if now.hour < 8:
        start_of_yesterday = datetime(now.year, now.month, now.day, 8) - timedelta(days=1)
    else:
        start_of_yesterday = datetime(now.year, now.month, now.day, 8)
    end_of_yesterday = start_of_yesterday + timedelta(days=1)

    # Get department statistics data
    departments_data, long_stay_patients = get_department_statistics(start_of_yesterday, end_of_yesterday, db)
    service_basis_data = get_service_basis_statistics(db)
    directions_data = get_directions_statistics(db)

    # Create a Word document
    doc = create_document(departments_data, long_stay_patients, service_basis_data, directions_data)

    # Create a temporary file
    temp_file_path = tempfile.NamedTemporaryFile(delete=False).name

    # Save the document to the temporary file
    doc.save(temp_file_path)

    # Return the temporary file to the client
    return FileResponse(temp_file_path, filename=f"Статистика на {datetime.now().strftime('%d.%m.%Y')}.docx")

def get_department_statistics(start_date, end_date, db: Session):
    now = datetime.now()
    if now.hour < 8:
        start_of_today = datetime(now.year, now.month, now.day, 0) 
    else:
        start_of_today = datetime(now.year, now.month, now.day, 0)
    end_of_today = start_of_today + timedelta(days=1)
    # Dictionary mapping departments to bed counts
    bed_counts = {
        'Хирургическое отделение': 50,
        'Терапевтическое отделение': 40,
        'Неврологическое отделение': 25,
        'Кожно-венерологическое отделение': 25,
        'Инфекционное отделение': 30
    }

    # Subquery for the latest patient movements
    latest_movements = (
        db.query(
            models.PatientMovement.patient_id,
            func.max(models.PatientMovement.event_date).label('latest_date')
        )
        .group_by(models.PatientMovement.patient_id)
        .subquery()
    )

    # Main query
    department_query = (
    db.query(
        models.PassportData.branch.label('department'),  # Используем branch как department
        func.count().filter(models.PatientMovement.event_type.in_([
                "Прием на госпитализацию",
                "Перевод в другое отделение",
                "Планируется перевод"
        ])).label('active_patients'),
        func.count().filter(and_(
            models.PatientMovement.event_type == 'Прием на госпитализацию',
            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today
        )).label('admitted'),
        func.count().filter(and_(
            or_(
                models.PatientMovement.event_type == 'Выписан',
                models.PatientMovement.event_type == 'Выписан за нарушение режима',
                models.PatientMovement.event_type == 'Направлен в санаторий',
            ),

            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today
        )).label('discharged'),
        func.count().filter(and_(
            models.PatientMovement.event_type == 'Перевод в другое отделение',
            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today,
        )).label('transferred_out'),
        func.count().filter(and_(
            models.PatientMovement.event_type == 'Перевод в другое отделение',
            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today
        )).label('transferred_in'),
        func.count().filter(and_(
            or_(
                models.PatientMovement.event_type == 'Перевод в другое МО',
                models.PatientMovement.event_type == 'Перевод в другое ВМО',
                models.PatientMovement.event_type == 'Перевод в РКБ',
            ),
            models.PatientMovement.event_date >= start_of_today,
            models.PatientMovement.event_date < end_of_today
        )).label('transferred_to_other_MO')
    )
    .join(models.PassportData, models.PatientMovement.patient_id == models.PassportData.id)  # Объединение с PassportData
    .join(latest_movements, and_(
        models.PatientMovement.patient_id == latest_movements.c.patient_id,
        models.PatientMovement.event_date == latest_movements.c.latest_date
    )) 
    .group_by(models.PassportData.branch))

    active_patients_query = (
        db.query(
            models.PatientMovement.department,
            func.count().label('active_patients')
        )
        .join(latest_movements, and_(
            models.PatientMovement.patient_id == latest_movements.c.patient_id,
            models.PatientMovement.event_date == latest_movements.c.latest_date
        ))
        .filter(models.PatientMovement.event_type == 'Прием на госпитализацию')
        .group_by(models.PatientMovement.department)
    )


    long_stay_patients_query = (
        db.query(
            models.PatientMovement.department.label('department'),
            models.PassportData.military_unit.label('military_unit'),
            models.HospitalData.diagnosis_upon_admission.label('diagnosis_upon_admission')
        )
        .join(models.PassportData, models.PatientMovement.patient_id == models.PassportData.id)
        .join(models.HospitalData, models.PatientMovement.patient_id == models.HospitalData.id)
        .filter(
            models.PatientMovement.event_type == 'Прием на госпитализацию',
            models.PatientMovement.event_date < end_date - timedelta(days=30)
        )
    )

    departments = department_query.all()
    long_stay_patients = long_stay_patients_query.all()


    # Assuming function to check if the diagnosis is surgical pathology
    def is_surgical_pathology(diagnosis):
        surgical_keywords = ['surgery', 'operation', 'surgical']
        return any(keyword in diagnosis.lower() for keyword in surgical_keywords)

    departments_data = [
        {
            "department": dept.department,
            "beds_deployed": bed_counts.get(dept.department, 0),
            "active_patients": dept.active_patients,
            "admitted": dept.admitted,
            "discharged": dept.discharged,
            "transferred_out": dept.transferred_out,
            "transferred_in": dept.transferred_in,
            "transferred_to_other_MO": dept.transferred_to_other_MO,
            "current_patients": dept.active_patients,
            "factual_number": dept.active_patients,
            "load_percentage": (dept.active_patients / bed_counts.get(dept.department, 1)) * 100,
            "free_beds": bed_counts.get(dept.department, 0) - dept.active_patients,
            "surgical_pathology": 0  # Placeholder, update as needed
        }
        for dept in departments
    ]

    return departments_data, long_stay_patients

def get_service_basis_statistics(db: Session):
    service_basis_query = (
        db.query(
            models.PassportData.service_basis,
            func.count(models.PassportData.id).label('count')
        )
        .group_by(models.PassportData.service_basis)
    )
    return service_basis_query.all()

def get_directions_statistics(db: Session):
    directions_query = (
        db.query(
            models.PassportData.directions,
            func.count(models.PassportData.id).label('count')
        )
        .group_by(models.PassportData.directions)
    )
    return directions_query.all()

@router.get("/generate_discharge_report/")
async def generate_discharge_report(request: Request, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    try:
        file_path = create_discharge_report(db)
        background_tasks.add_task(delete_file_with_delay, file_path, 300, background_tasks)
        return FileResponse(path=file_path, filename=os.path.basename(file_path), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": download_link})
    except Exception as e:
        logger.error(f"Error generating discharge report: {e}")
        return JSONResponse(status_code=500, content={"detail": f"Internal server error: {str(e)}"})

@router.get("/generate_discharge_report_v2/")
async def generate_discharge_report(request: Request, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    try:
        file_path = create_discharge_report_v2(db)
        background_tasks.add_task(delete_file_with_delay, file_path, 300, background_tasks)
        return FileResponse(path=file_path, filename=os.path.basename(file_path), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return JSONResponse(content={"message": "Report generated successfully.", "download_link": download_link})
    except Exception as e:
        logger.error(f"Error generating discharge report: {e}")
        return JSONResponse(status_code=500, content={"detail": f"Internal server error: {str(e)}"})



#region vvk
@router.get("/generate_vvk_report/")
def generate_vvk_report(date: datetime, db: Session = Depends(get_db)):
    vvk_records = db.query(
        models.VvkHospital,
        models.PassportData.full_name,
        models.PassportData.birthday_date,
        models.PassportData.military_rank,
        models.PassportData.military_unit,
        models.PassportData.personal_data,
        models.HospitalData.expert_diagnosis,
        models.VvkHospital.conclusion
    ).join(
        models.PassportData, models.VvkHospital.patient_id == models.PassportData.id
    ).outerjoin(
        models.HospitalData, models.PassportData.id == models.HospitalData.patient_id
    ).filter(
        models.VvkHospital.vvk_date == date
    ).all()

    if not vvk_records:
        raise HTTPException(status_code=404, detail="No VVK records found for this date")

    document = Document()
    section = document.sections[0]

    # Set page width and height for landscape orientation
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Add header
    header = document.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    header_run = header.add_run("Протокол заседания ВВК №19 ВГ1 от 02 февраля 2024 года.\n")
    header_run.bold = True
    set_font_and_spacing(header)

    header_run = header.add_run("Военный Госпиталь (На 150 Коек, Г. Казань)\n")
    set_font_and_spacing(header)

    header_run = header.add_run("Федеральное Государственное Казенное Учреждение «354 Военный Клинический   Госпиталь» Министерства Обороны Российской Федерации (ВГ (На 150 Коек, Г. Казань)) г. Казань, ул. Рауиса Гареева, д.120, 420064.\n")
    set_font_and_spacing(header)

    # Add table
    table = document.add_table(rows=1, cols=8)
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    headers = ['№п/п', 'ФИО', 'Дата рождения', 'Воинское звание', 'Воинская часть', 'Личный номер', 'Диагноз', 'Заключение ВВК']
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        set_font_and_spacing(hdr_cells[i].paragraphs[0])

    for idx, record in enumerate(vvk_records, start=1):
        row_cells = table.add_row().cells
        data = [
            str(idx),
            record.full_name,
            record.birthday_date.strftime('%d.%m.%Y') if record.birthday_date else '',
            record.military_rank,
            record.military_unit,
            record.personal_data,
            record.expert_diagnosis if record.expert_diagnosis else '',
            record.conclusion
        ]
        for i, text in enumerate(data):
            row_cells[i].text = text
            set_font_and_spacing(row_cells[i].paragraphs[0])

    # Add footer lines
    footer_lines = [
        "\nЧлены комиссии:",
        "Председатель ВВК: ________________________________________________________________________________________________________________",
        "Член комиссии: ___________________________________________________________________________________________________________________",
        "Член комиссии: ___________________________________________________________________________________________________________________",
        "Член комиссии: ___________________________________________________________________________________________________________________",
        "Член комиссии: ___________________________________________________________________________________________________________________",
        "Секретарь ВВК: ___________________________________________________________________________________________________________________"
    ]
    for line in footer_lines:
        footer_paragraph = document.add_paragraph(line)
        set_font_and_spacing(footer_paragraph)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    # Encode filename for Content-Disposition header
    filename = f"Протокол ВВК на {date.strftime('%d.%m.%Y')}.docx"
    encoded_filename = urllib.parse.quote(filename)

    return StreamingResponse(
        file_stream,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )
#endregion

#endregion

#region Допы

@router.get("/in_work", response_class=HTMLResponse)
async def read_root(request: Request):
    template = "index.html"
    return templates.TemplateResponse(template, {"request": request})

@router.get("/instruction", response_class=HTMLResponse)
async def read_root(request: Request):
    template = "instruction.html"
    return templates.TemplateResponse(template, {"request": request})

@router.get("/history_of_discovery", response_class=HTMLResponse)
async def read_root(request: Request):
    template = "about/history_of_discovery.html"
    return templates.TemplateResponse(template, {"request": request})

@router.get("/documents", response_class=HTMLResponse)
async def read_root(request: Request):
    template = "documents.html"
    return templates.TemplateResponse(template, {"request": request})

#endregion