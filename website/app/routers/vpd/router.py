from typing import List
from fastapi import APIRouter, Depends, HTTPException, Request, Form
from sqlalchemy.orm import Session
from sqlalchemy import create_engine, func, select
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from sqlalchemy.orm import sessionmaker
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse
from ...model.vpd.models import VPD  # Импортируем модель VPD
from ...model.vpd.schemas import VPDSchema
from ...model.database.database import SessionLocal
from .._utils.documents_utils import fill_template_v3, docs_dir, templates
import os
from .._utils.read_excel_utils import read_excel, read_excel_ds
router = APIRouter()
current_dir = os.path.dirname(os.path.abspath(__file__))

templates = Jinja2Templates(directory=os.path.abspath(os.path.join(current_dir, "../../templates")))



def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

from docx import Document
from datetime import date
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import locale

@router.get("/generate-report/")
async def generate_report(selected_date: date, db: Session = Depends(get_db)):
    # Установка локали для корректного отображения даты на русском языке
    


    records = db.query(VPD).filter(VPD.departure_date == selected_date).all()
    
    formatted_date = selected_date.strftime(f'«%d» %B %Y г.')

  

    formatted_date = selected_date.strftime(f'«%d» %B %Y г.')


    if not records:
        raise HTTPException(status_code=404, detail="No records found for the selected date")

    doc = Document()

        # Настройка полей страницы
    section = doc.sections[0]
    section.top_margin = Pt(42.50)  # 1 см в пунктах (1 см = 28.35 пунктов)
    section.bottom_margin = Pt(42.85)  # 1.5 см в пунктах (например)
    section.left_margin = Pt(42.85)  # 1 см в пунктах
    section.right_margin = Pt(42.85)  # 1 см в пунктах

    # Настройка шрифта для всего документа
    styles = doc.styles
    normal_style = styles['Normal']
    font = normal_style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_bold_centered_paragraph(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.bold = True  # Жирный шрифт
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.line_spacing = Pt(12)  # Межстрочный интервал 1

        p.paragraph_format.space_after = Pt(0)  # Расстояние после параграфа 1 пункт

    def add_centered_paragraph(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.line_spacing = Pt(12)  # Межстрочный интервал 1
        p.paragraph_format.space_after = Pt(0)  # Расстояние после параграфа 1 пункт

    # Функция для перевода месяцев на русский язык
    def month_to_russian(month_number):
        months = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря"
        ]
        return months[month_number - 1]


    # Дата на русском языке
    day = selected_date.day
    month = month_to_russian(selected_date.month)
    year = selected_date.year
    formatted_date = f'«{day}» {month} {year} г.'

    for record in records:
        # Основной контент


        add_bold_centered_paragraph('ВЫПИСКА ИЗ ПРИКАЗА')
        add_bold_centered_paragraph('НАЧАЛЬНИКА ВОЕННОГО ГОСПИТАЛЯ (на 150 коек, г. Казань)')
        add_bold_centered_paragraph('ФГКУ «354 ВКГ» Минобороны России')
        add_centered_paragraph('по строевой части')
        add_centered_paragraph('№ 99')
        doc.add_paragraph(f'{formatted_date}                                                                                           г. Казань')
        p = doc.add_paragraph()
        
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.line_spacing = Pt(12)  # Межстрочный интервал 1
        p.paragraph_format.space_after = Pt(0)  # Расстояние после параграфа 1 пункт

        doc.add_paragraph(
            f'            В связи с окончанием лечения снять с котлового довольствия {record.military_rank}, {record.full_name} {record.departure_date.strftime("%d.%m.%Y")} г.р., с ужина {formatted_date} года. Полагать убывшим в {record.where_to}.'
        )
        doc.add_paragraph(' На путь следования выдать:')
        p = doc.add_paragraph()
        p.add_run('-').bold = True  # Жирный символ
        p.add_run(' Предписание.')

        doc.add_paragraph('            Основание: указания начальника Генерального Штаба МО РФ от 06.12.2022 г. №308/1/3867дсп, указание ГШ ВС РФ от 20.05.2023 года №314/2/4897 дсп, рапорт военнослужащего.')



        p = doc.add_paragraph('Начальник ВГ (на 150 коек, г. Казань) ФГКУ «354 ВКГ» Минобороны России', style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Центрирование
        p.paragraph_format.line_spacing = Pt(12)  # Межстрочный интервал 1
        run = p.runs[0]
        run.bold = True  # Жирный шрифт
        
        p = doc.add_paragraph('подполковник медицинской службы                                                  Д.Ижбердеев', style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Центрирование
        p.paragraph_format.line_spacing = Pt(12)  # Межстрочный интервал 1
        run = p.runs[0]
        run.bold = True  # Жирный шрифт
        
        doc.add_paragraph('')
        p = doc.add_paragraph('Выписка верна: архивариус                                                    А. Шайхразыева', style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Центрирование
        p.paragraph_format.line_spacing = Pt(12)  # Межстрочный интервал 1
        run = p.runs[0]
        run.bold = True  # Жирный шрифт

    file_path = f"Выписка из приказа {selected_date}.docx"
    doc.save(file_path)

    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=file_path)



@router.get("/generate_prescription/{patient_id}", response_class=FileResponse)
async def generate_prescription(patient_id: int, db: Session = Depends(get_db)):
    # Получаем данные пациента из базы данных
    patient = db.query(VPD).filter(VPD.id == patient_id).first()
    
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    # Формируем словарь с данными для заполнения шаблона
    patient_data = {
        "Звание": patient.military_rank,
        "ФИО": patient.full_name,
        "Дата убытия": f"{patient.departure_date.strftime('%d.%m.%Y')} г.",
        "Убыть куда": patient.where_to,
        "Дата прибытия": f"{patient.arrival_date.strftime('%d.%m.%Y')} г.",
    }

    # Путь к шаблону документа
    template_path = os.path.join(docs_dir, "templates", "Предписание, вариант с таблицами.docx")

    # Генерация документа (необходимая функция fill_template)
    output_path = fill_template_v3(template_path, patient_data, patient.full_name, file_name=f'Преписание {patient.where_to}')

    # Возвращаем файл для скачивания
    return FileResponse(output_path, filename=os.path.basename(output_path), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')




from sqlalchemy import and_
from fastapi.responses import JSONResponse

@router.post("/edit-vpd/")
async def edit_vpd(request: Request):
    form = await request.form()
    id = form.get('id')
    military_rank = form.get('military_rank')
    full_name = form.get('full_name')
    military_unit = form.get('military_unit')
    service_base = form.get('service_base')
    departure_date = form.get('departure_date')
    where_to = form.get('where_to')
    arrival_date = form.get('arrival_date')
    current_time = form.get('current_time')

    if not id:
        raise HTTPException(status_code=400, detail="ID is required")

    db: Session = SessionLocal()
    vpd_record = db.query(VPD).filter(VPD.id == id).first()

    if not vpd_record:
        db.close()
        raise HTTPException(status_code=404, detail="Record not found")

    # Update record fields
    vpd_record.military_rank = military_rank
    vpd_record.full_name = full_name
    vpd_record.military_unit = military_unit
    vpd_record.service_base = service_base
    vpd_record.departure_date = departure_date
    vpd_record.where_to = where_to
    vpd_record.arrival_date = arrival_date
    vpd_record.current_time = current_time

    db.commit()
    db.refresh(vpd_record)
    db.close()

    return JSONResponse(content={"success": True}, status_code=200)

@router.get("/search/")
async def search_vpd(query: str, db: Session = Depends(get_db)):
    results = db.query(VPD).filter(
        (VPD.full_name.ilike(f'%{query}%')) |
        (VPD.military_unit.ilike(f'%{query}%'))
    ).all()

    if not results:
        return {"message": "No records found"}
    
    # Преобразуем результаты в список словарей для возврата в формате JSON
    return [
        {
            "id": record.id,
            "full_name": record.full_name,
            "military_unit": record.military_unit,
            "departure_date": record.departure_date.strftime("%d.%m.%Y"),
            "where_to": record.where_to,
            "arrival_date": record.arrival_date.strftime("%d.%m.%Y"),
            "military_rank": record.military_rank,
            "service_base": record.service_base,
            "current_time": record.current_time
        } 
        for record in results
    ]


@router.get("/", response_class=HTMLResponse)
async def read_vpd(request: Request, db: Session = Depends(get_db), full_name: str = None, date: str = None):
    query = db.query(VPD)
    
    if full_name:
        query = query.filter(VPD.full_name.ilike(f"%{full_name}%"))
    
    if date:
        # Допустим, date приходит в формате 'YYYY-MM-DD'
        query = query.filter(func.date(VPD.departure_date) == date)
    
    vpd_records = query.all()
    
    return templates.TemplateResponse("vpd/vpd.html", {"request": request, "vpd_records": vpd_records})

# Отдельный API-эндпоинт для асинхронного поиска
@router.get("/search", response_model=List[VPDSchema])
async def search_vpd(full_name: str = None, date: str = None, db: Session = Depends(get_db)):
    query = db.query(VPD)
    
    if full_name:
        query = query.filter(VPD.full_name.ilike(f"%{full_name}%"))
    
    if date:
        query = query.filter(func.date(VPD.departure_date) == date)
    
    vpd_records = query.all()
    
    return vpd_records

from datetime import date, datetime

@router.get("/autocomplete/")
async def autocomplete(query: str, db: Session = Depends(get_db)):
    results = db.query(VPD.where_to).filter(VPD.military_unit.ilike(f"%{query}%")).distinct().all()
    suggestions = [result.where_to for result in results]
    print(results)
    print(query)
    return suggestions

@router.get("/get-data")
async def get_data(db: Session = Depends(get_db)):
    records = db.query(VPD).all()
    return [
        {
            "id": record.id,
            "full_name": record.full_name,
            "military_unit": record.military_unit,
            "departure_date": record.departure_date.strftime("%d.%m.%Y"),
            "where_to": record.where_to,
            "arrival_date": record.arrival_date.strftime("%d.%m.%Y"),
            "military_rank": record.military_rank,
            "service_base": record.service_base,
            "current_time": record.current_time.strftime("%d.%m.%Y")
        }
        for record in records
    ]

from io import BytesIO
from openpyxl import Workbook
from urllib.parse import quote

@router.get("/export")
def export_vpd_to_excel(db: Session = Depends(get_db)):
    # Получаем все записи из базы данных
    records = db.query(VPD).all()
    
    if not records:
        raise HTTPException(status_code=404, detail="No records found")
    
    # Создаем Excel-файл
    wb = Workbook()
    ws = wb.active
    today = datetime.now().strftime('%d_%m_%Y')
    ws.title = f"ВПД на {today}"

    # Добавляем заголовки
    headers = [
        "ID", 
        "Воинское звание", 
        "ФИО", 
        "Дата убытия", 
        "Куда направляется", 
        "Дата прибытия", 
        "Воинская часть", 
        "Основа службы", 
        "Время создания предписания"
    ]
    ws.append(headers)
    
    for record in records:
        ws.append([
            record.id,
            record.military_rank,
            record.full_name,
            record.departure_date.strftime('%d.%m.%Y') if record.departure_date else None,
            record.where_to,
            record.arrival_date.strftime('%d.%m.%Y') if record.arrival_date else None,
            record.military_unit,
            record.service_base,
            record.current_time.strftime('%d.%m.%Y') if record.current_time else None,
        ])

    # Сохраняем файл в буфер
    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    

    file_name = f"ВПД на {today}.xlsx"
    encoded_file_name = quote(file_name.encode('utf-8'))
    
    return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f"attachment; filename={encoded_file_name}"})

@router.delete("/delete-vpd/{record_id}")
async def delete_vpd(record_id: int):
    db: Session = SessionLocal()
    vpd_record = db.query(VPD).filter(VPD.id == record_id).first()

    if not vpd_record:
        db.close()
        raise HTTPException(status_code=404, detail="Record not found")

    db.delete(vpd_record)
    db.commit()
    db.close()

    return JSONResponse(content={"success": True}, status_code=200)

@router.post("/add-vpd/")
async def add_vpd(
    military_rank: str = Form(...),
    full_name: str = Form(...),
    departure_date: date = Form(...),
    where_to: str = Form(...),
    arrival_date: date = Form(...),
    military_unit: str = Form(...),
    service_base: str = Form(...),
    db: Session = Depends(get_db)
):

    vpd_data = VPDSchema(
        military_rank=military_rank,
        full_name=full_name,
        departure_date=departure_date,
        where_to=where_to,
        arrival_date=arrival_date,
        military_unit=military_unit,
        service_base=service_base,
        current_time=datetime.now().date()
    )

    vpd = VPD(**vpd_data.dict())
    db.add(vpd)
    db.commit()
    return {"message": "Record added successfully"}


