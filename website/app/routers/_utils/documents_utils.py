from fastapi import APIRouter, HTTPException, Depends, Request, BackgroundTasks
from fastapi.templating import Jinja2Templates
import logging
from datetime import datetime
import time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

import os

current_dir = os.path.dirname(os.path.abspath(__file__))

docs_dir = os.path.abspath(os.path.join(current_dir, "../../static/docs"))

current_dir = os.path.dirname(os.path.abspath(__file__))

img_dir = os.path.abspath(os.path.join(current_dir, "../static/img"))
js_dir = os.path.abspath(os.path.join(current_dir, "../static/js"))
css_dir = os.path.abspath(os.path.join(current_dir, "../static/css"))
templates = Jinja2Templates(directory=os.path.abspath(os.path.join(current_dir, "../../templates")))

logger = logging.getLogger(__name__)

def fill_template(template_path: str, context: dict, full_name: str, file_name: str = 'Справка о нарушении') -> str:
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        modified_text = full_text
        for key, value in context.items():
            placeholder = '{{' + key + '}}'
            if placeholder in full_text:
                modified_text = modified_text.replace(placeholder, str(value))
        
        if modified_text != full_text:
            for run in paragraph.runs:
                run.clear()
            new_run = paragraph.add_run(modified_text)
            if paragraph.runs:
                new_run.bold = paragraph.runs[0].bold
                new_run.italic = paragraph.runs[0].italic
                new_run.underline = paragraph.runs[0].underline
                new_run.font.size = paragraph.runs[0].font.size
                new_run.font.color.rgb = paragraph.runs[0].font.color.rgb

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(run.text for run in paragraph.runs)
                    modified_text = full_text
                    for key, value in context.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in full_text:
                            modified_text = modified_text.replace(placeholder, str(value))
                    
                    if modified_text != full_text:
                        for run in paragraph.runs:
                            run.clear()
                        new_run = paragraph.add_run(modified_text)
                        if paragraph.runs:
                            new_run.bold = paragraph.runs[0].bold
                            new_run.italic = paragraph.runs[0].italic
                            new_run.underline = paragraph.runs[0].underline
                            new_run.font.size = paragraph.runs[0].font.size
                            new_run.font.color.rgb = paragraph.runs[0].font.color.rgb
    
    def format_name(full_name):
        parts = full_name.split()
        if len(parts) == 3:
            formatted_name = f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
        elif len(parts) == 2:
            formatted_name = f"{parts[0]} {parts[1][0]}."
        else:
            formatted_name = full_name
        return formatted_name

    if not os.path.exists(docs_dir):
        os.makedirs(docs_dir)

    output_path = os.path.join(docs_dir, f'{file_name} {format_name(full_name)} {datetime.now().date()}.docx')
    doc.save(output_path)
    return output_path

def fill_template_v3(template_path: str, context: dict, full_name: str, file_name: str = 'Справка о нарушении') -> str:
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        modified_text = full_text
        for key, value in context.items():
            placeholder = '{{' + key + '}}'
            if placeholder in full_text:
                modified_text = modified_text.replace(placeholder, str(value))
        
        if modified_text != full_text:
            for run in paragraph.runs:
                run.clear()
            new_run = paragraph.add_run(modified_text)
            if paragraph.runs:
                new_run.bold = paragraph.runs[0].bold
                new_run.italic = paragraph.runs[0].italic
                new_run.underline = paragraph.runs[0].underline
                new_run.font.size = paragraph.runs[0].font.size
                new_run.font.color.rgb = paragraph.runs[0].font.color.rgb

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(run.text for run in paragraph.runs)
                    modified_text = full_text
                    for key, value in context.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in full_text:
                            modified_text = modified_text.replace(placeholder, str(value))
                    
                    if modified_text != full_text:
                        for run in paragraph.runs:
                            run.clear()
                        new_run = paragraph.add_run(modified_text)
                        if paragraph.runs:
                            new_run.bold = paragraph.runs[0].bold
                            new_run.italic = paragraph.runs[0].italic
                            new_run.underline = paragraph.runs[0].underline
                            new_run.font.size = paragraph.runs[0].font.size
                            new_run.font.color.rgb = paragraph.runs[0].font.color.rgb
    


    if not os.path.exists(docs_dir):
        os.makedirs(docs_dir)

    output_path = os.path.join(docs_dir, f'{file_name} {full_name} {datetime.now().date()}.docx')
    doc.save(output_path)
    return output_path


def set_paragraph_formatting(paragraph, space_before=Pt(0), space_after=Pt(0), line_spacing=1.0):
    """
    Устанавливает форматирование абзаца.

    Параметры:
    paragraph (docx.text.paragraph.Paragraph): Абзац, к которому применяется форматирование.
    space_before (Pt, optional): Пробел перед абзацем. По умолчанию 0 пунктов.
    space_after (Pt, optional): Пробел после абзаца. По умолчанию 0 пунктов.
    line_spacing (float, optional): Междустрочный интервал. По умолчанию 1.0.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = space_before
    paragraph_format.space_after = space_after
    paragraph_format.line_spacing = line_spacing

def add_custom_run(paragraph, text, bold=False, size=Pt(11), underline=False):
    """
    Добавляет текст в абзац с настройкой шрифта и возможностью подчеркивания.

    Параметры:
    paragraph (docx.text.paragraph.Paragraph): Абзац, к которому добавляется текст.
    text (str): Текст, который нужно добавить.
    bold (bool, optional): Сделать текст жирным. По умолчанию False.
    size (Pt, optional): Размер шрифта. По умолчанию 11 пунктов.
    underline (bool, optional): Подчеркнуть текст. По умолчанию False.
    """
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    run.underline = underline

def add_custom_run_yellow(paragraph, text, bold=False, size=Pt(11), highlight=None):
    """
    Добавляет текст в абзац с настройкой шрифта и возможностью подсветки.

    Параметры:
    paragraph (docx.text.paragraph.Paragraph): Абзац, к которому добавляется текст.
    text (str): Текст, который нужно добавить.
    bold (bool, optional): Сделать текст жирным. По умолчанию False.
    size (Pt, optional): Размер шрифта. По умолчанию 11 пунктов.
    highlight (docx.oxml.enum.WdColor, optional): Цвет подсветки текста. По умолчанию None.
    """
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    if highlight:
        run.font.highlight_color = highlight

def set_cell_border(cell, border_color="000000", border_size=1):
    """
    Устанавливает пользовательские границы для ячейки таблицы.

    Параметры:
    cell (docx.table.Cell): Ячейка таблицы, для которой устанавливаются границы.
    border_color (str, optional): Цвет границы в формате HEX. По умолчанию "000000".
    border_size (int, optional): Толщина границы в пунктах. По умолчанию 1.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size * 8))
        border.set(qn('w:color'), border_color)
        tcPr.append(border)

def set_cell_vertical_alignment(cell, align="center"):
    """
    Устанавливает вертикальное выравнивание для ячейки таблицы.

    Параметры:
    cell (docx.table.Cell): Ячейка таблицы, для которой устанавливается выравнивание.
    align (str, optional): Вертикальное выравнивание. Может быть "top", "center" или "bottom". По умолчанию "center".
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcValign = OxmlElement('w:vAlign')
    tcValign.set(qn('w:val'), align)
    tcPr.append(tcValign)

def create_header_with_doctor_and_shape(doc, doctor_name, shape_text, shape_width_inches):
    """
    Создает заголовок с именем врача и фигурой в таблице.

    Параметры:
    doc (docx.Document): Документ, к которому добавляется таблица.
    doctor_name (str): Имя лечащего врача.
    shape_text (str): Текст для отображения в фигуре.
    shape_width_inches (float): Ширина фигуры в дюймах.
    """
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    
    # Настройка столбцов таблицы
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 2)
    left_cell.width = Inches(shape_width_inches)
    right_cell.width = Inches(3)  # Ширина столбца с текстом врача
    
    # Добавление текста с указанием лечащего врача
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_custom_run(right_paragraph, f"Лечащий врач: {doctor_name}\n", bold=True)

    # Создание фигуры с текстом в левой ячейке
    left_cell.text = shape_text
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = left_paragraph.runs[0]
    run.font.bold = True
    run.font.size = Pt(8)
    set_cell_border(left_cell, border_color="000000", border_size=1)
    set_cell_vertical_alignment(left_cell, 'center')

    # Настройка ширины ячеек
    left_cell.width = Inches(shape_width_inches)
    table.width = left_cell.width + right_cell.width

def center_text_in_line(text, line_width=68):
    """
    Центрирует текст в строке заданной ширины, заполняя пробелами до и после текста.
    
    :param text: Текст для центрирования
    :param line_width: Ширина строки в символах
    :return: Строка с центрированным текстом
    """
    text_length = len(text)
    if text_length >= line_width:
        return text[:line_width]  # Обрезать текст, если он слишком длинный

    total_padding = line_width - text_length
    left_padding = total_padding // 2
    right_padding = total_padding - left_padding

    # Собираем итоговую строку с левым отступом, текстом и правым отступом
    return ' ' * left_padding + text + ' ' * right_padding

def add_full_underlined_text(doc, main_text, superscript_text, font_size=Pt(12)):
    """
    Добавляет подчеркивающий текст с подстрочным текстом в документ Word.

    Параметры:
    doc (docx.Document): Объект документа, в который добавляется текст.
    main_text (str): Основной текст, который будет подчеркиваться.
    superscript_text (str): Подстрочный текст, который добавляется после основного текста.
    font_size (Pt): Размер шрифта для основного текста (по умолчанию 12 пунктов).

    Описание:
    Функция добавляет в документ абзац с основным подчеркивающимся текстом, 
    центрирует его и добавляет подстрочный текст под ним. Размер шрифта основного 
    текста можно настроить через параметр `font_size`.
    """
    section = doc.sections[0]
    page_width_pt = section.page_width.pt  # Ширина страницы в пунктах
    left_margin_pt = section.left_margin.pt
    right_margin_pt = section.right_margin.pt
    usable_width_pt = page_width_pt - (left_margin_pt + right_margin_pt)

    approx_char_width = font_size / 2  # Приблизительная ширина символа в пунктах
    max_chars = int(usable_width_pt / approx_char_width)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE
    
    text_length = len(main_text)
    total_padding = max_chars - text_length
    padding_each_side = total_padding // 2

    run = paragraph.add_run(' ' * padding_each_side)
    run.underline = True
    run.font.size = font_size

    run = paragraph.add_run(main_text)
    run.font.name = 'Times New Roman'
    run.font.size = font_size
    run.underline = True

    run = paragraph.add_run(' ' * (total_padding - padding_each_side))
    run.underline = True
    run.font.size = font_size

    run = paragraph.add_run('\n' + superscript_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.font.superscript = True

def set_font_to_times_new_roman(run):
    """
    Устанавливает шрифт для текста в объекте Run на 'Times New Roman'.

    Параметры:
    run (docx.text.run.Run): Объект Run, для которого устанавливается шрифт.
    """
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def fill_template_v_times(template_path: str, context: dict, full_name, file_name='Справка о ранении') -> str:
    """
    Заполняет шаблон документа Word данными из контекста и сохраняет его с именем файла.

    Параметры:
    template_path (str): Путь к файлу шаблона документа Word.
    context (dict): Словарь с данными для замены в шаблоне.
    full_name (str): Полное имя для включения в документ.
    file_name (str): Имя файла для сохранения документа (по умолчанию 'Справка о ранении').

    Возвращает:
    str: Путь к сохраненному документу.
    """
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        modified_text = full_text
        for key, value in context.items():
            placeholder = '{{' + key + '}}'
            if placeholder in full_text:
                modified_text = modified_text.replace(placeholder, str(value))
        
        if modified_text != full_text:
            # Удаляем старый текст
            for run in paragraph.runs:
                run.clear()
            # Добавляем измененный текст с сохранением форматирования
            new_run = paragraph.add_run(modified_text)
            # Копируем стиль первоначального run, если он был
            if paragraph.runs:
                new_run.bold = paragraph.runs[0].bold
                new_run.italic = paragraph.runs[0].italic
                new_run.underline = paragraph.runs[0].underline
                new_run.font.size = paragraph.runs[0].font.size
                new_run.font.color.rgb = paragraph.runs[0].font.color.rgb
            # Устанавливаем шрифт Times New Roman
            set_font_to_times_new_roman(new_run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(run.text for run in paragraph.runs)
                    modified_text = full_text
                    for key, value in context.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in full_text:
                            modified_text = modified_text.replace(placeholder, str(value))
                    
                    if modified_text != full_text:
                        # Очистка существующих runs
                        for run in paragraph.runs:
                            run.clear()
                        # Добавление измененного текста с сохранением форматирования
                        new_run = paragraph.add_run(modified_text)
                        # Копирование стиля из оригинального run, если есть
                        if paragraph.runs:
                            new_run.bold = paragraph.runs[0].bold
                            new_run.italic = paragraph.runs[0].italic
                            new_run.underline = paragraph.runs[0].underline
                            new_run.font.size = paragraph.runs[0].font.size
                            new_run.font.color.rgb = paragraph.runs[0].font.color.rgb
                        # Установка шрифта на Times New Roman
                        set_font_to_times_new_roman(new_run)
    


    
    output_path = os.path.join(docs_dir, f'{file_name} {format_name(full_name)} {datetime.now().date()}.docx')
    doc.save(output_path)
    return output_path

def format_name(full_name):
    """
    Форматирует полное имя в виде "Фамилия И.О.".

    Параметры:
    full_name (str): Полное имя.

    Возвращает:
    str: Отформатированное имя.
    """
    parts = full_name.split()
    if len(parts) == 3:
        formatted_name = f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    elif len(parts) == 2:
        formatted_name = f"{parts[0]} {parts[1][0]}."
    else:
        formatted_name = full_name
    return formatted_name

def delete_file_with_delay(file_path: str, delay: int, background_tasks: BackgroundTasks):
    """
    Запускает фоновую задачу для удаления файла после заданной задержки.

    Параметры:
    file_path (str): Путь к файлу, который нужно удалить.
    delay (int): Задержка перед удалением файла в секундах.
    background_tasks (BackgroundTasks): Объект для управления фоновыми задачами.

    Описание:
    Функция создает задачу, которая удаляет указанный файл после указанной задержки. Задача добавляется 
    в очередь фоновых задач и выполняется асинхронно. При успешном удалении файл будет удален, и в 
    журнале будет зафиксировано сообщение об успешном удалении. В случае ошибки удаления будет записано 
    сообщение об ошибке.
    """
    def delayed_deletion():
        time.sleep(delay)  # Задержка в секундах
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"File {file_path} has been successfully deleted.")
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")

    # Добавляем задачу в фоновый процесс
    background_tasks.add_task(delayed_deletion)

def set_cell_border_v2(cell):
    """
    Устанавливает черные границы для ячейки таблицы.

    Параметры:
    cell (docx.table.Cell): Ячейка таблицы, для которой устанавливаются границы.

    Описание:
    Функция устанавливает границы ячейки таблицы с использованием черного цвета. Границы будут установлены 
    по всем сторонам (верх, низ, слева, справа) и внутри ячейки. Толщина границы задается в 1 пункт.
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        element = tcPr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcPr.append(element)
            
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')  # 1/8 pt, 8 * 0.125 = 1 pt
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')

def set_table_borders(table):
    """
    Устанавливает черные границы для всей таблицы.

    Параметры:
    table (docx.table.Table): Таблица, для которой устанавливаются границы.

    Описание:
    Функция устанавливает границы для всей таблицы, включая внешние и внутренние границы ячеек. 
    Все границы имеют черный цвет, толщину 1 пункт и нулевое пространство между границами.
    """
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

def set_font_and_spacing(paragraph):
    """
    Устанавливает шрифт и межстрочный интервал для абзаца.

    Параметры:
    paragraph (docx.text.paragraph.Paragraph): Абзац, для которого устанавливаются шрифт и интервал.

    Описание:
    Функция устанавливает шрифт 'Times New Roman' и размер шрифта 10 пунктов для всех runs (участков текста) 
    в абзаце. Межстрочный интервал устанавливается в 12 пунктов.
    """
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    paragraph.paragraph_format.line_spacing = Pt(12)  # single line spacing in pt
